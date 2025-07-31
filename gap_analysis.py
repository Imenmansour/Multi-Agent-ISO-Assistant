import streamlit as st
import os
import json
from typing import List
from langchain_qdrant import QdrantVectorStore
from sentence_transformers import SentenceTransformer
from langchain.embeddings.base import Embeddings
import json
import re
import pandas as pd
import os
from docx import Document as DocxDocument
from docx.shared import RGBColor
import torch
from transformers import AutoModelForSequenceClassification, AutoTokenizer
from langchain_ollama import ChatOllama
from transformers import AutoTokenizer
from langchain_ollama import ChatOllama
from transformers import AutoTokenizer
from transformers import AutoTokenizer
from typing import List
from node.company_docs import ask_company_question
print("Setting up graphs and language models...")

# Define the database file path for chat history
DB_FILE = "db.json"



# -----------------------------
# Initialize Language Models
# -----------------------------
print("Initializing language models...")



llm_ollama =ChatOllama(model="qwen3:0.6b",temperature=0.7)

gemma_model=ChatOllama(model="qwen3:0.6b",temperature=0.7)
model_id=r"C:\Users\G800613RTS\bge-reranker-v2-m3"
reranker_tokenizer = AutoTokenizer.from_pretrained(model_id)
reranker_model = AutoModelForSequenceClassification.from_pretrained(model_id)
reranker_model.eval()
#reranker 
def rerank_documents(query, documents, top_k=2):
    doc_texts = []
    docs_list = []
    
    for doc in documents:
        # Handle tuple from similarity_search_with_score
        if isinstance(doc, tuple) and len(doc) == 2:
            document = doc[0]  # Extract the document from the (doc, score) tuple
            docs_list.append(document)
            
            if hasattr(document, 'page_content'):
                doc_texts.append(document.page_content)
            elif isinstance(document, dict) and 'text' in document:
                doc_texts.append(document['text'])
            elif isinstance(document, dict) and 'chunk' in document:
                doc_texts.append(document['chunk'])
            elif isinstance(document, dict) and 'summary_chunk' in document:
                doc_texts.append(document['summary_chunk'])
            elif isinstance(document, dict) and 'content' in document:
                doc_texts.append(document['content'])
            elif isinstance(document, str):
                doc_texts.append(document)
            else:
                print(f"Warning: Skipping unsupported document type in tuple: {type(document)}")
                continue
        # Original handling for direct documents
        elif hasattr(doc, 'page_content'):
            doc_texts.append(doc.page_content)
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'text' in doc:
            doc_texts.append(doc['text'])
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'chunk' in doc:
            doc_texts.append(doc['chunk'])
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'summary_chunk' in doc:
            doc_texts.append(doc['summary_chunk'])
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'content' in doc:
            doc_texts.append(doc['content'])
            docs_list.append(doc)
        elif isinstance(doc, str):
            doc_texts.append(doc)
            docs_list.append(doc)
        else:
            print(f"Warning: Skipping unsupported document type: {type(doc)}")
            continue

    pairs = [[query, text] for text in doc_texts]

    with torch.no_grad():
        inputs = reranker_tokenizer(pairs, padding=True, truncation=True, return_tensors='pt', max_length=512)
        scores = reranker_model(**inputs, return_dict=True).logits.view(-1, ).float()

    scored_docs = list(zip(scores.tolist(), docs_list))
    scored_docs.sort(key=lambda x: x[0], reverse=True)

    return [doc for _, doc in scored_docs[:top_k]]
# -----------------------------
# Reload embeddings model
# -----------------------------
print("Loading embedding model...")
model = SentenceTransformer(
    "nomic-ai/nomic-embed-text-v1",
    trust_remote_code=True
)

class CustomSentenceTransformerEmbeddings(Embeddings):  # Inherit from LangChain base
    def __init__(self, model):
        self.model = model

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self.model.encode(texts, normalize_embeddings=True).tolist()

    def embed_query(self, text: str) -> List[float]:
        return self.model.encode([text], normalize_embeddings=True)[0].tolist()

embeddings = CustomSentenceTransformerEmbeddings(model)

# -----------------------------
# Connect to Vector Stores
# -----------------------------
vector_store_i_title = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="io_title_collection",
    url="http://localhost:6333",
)
vector_store_i_text = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="io_text_collection",
    url="http://localhost:6333",
)
vector_store_summaries= QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="co_summary_vectors",
    url="http://localhost:6333",
)
vector_store_originals = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="co_original_vectors",
    url="http://localhost:6333",
)
# -----------------------------
# Gap Analysis Logic
# -----------------------------
print("Setting up gap analysis logic...")
# Load theme clauses
with open("sheet_test.json", "r",encoding="utf-8") as f:
    theme_clauses = json.load(f)


def save_to_csv(results, filename="gap_analysis_results.csv"):
    import pandas as pd
    df = pd.DataFrame(results)
    
    # Save to the current directory with absolute path
    import os
    csv_path = os.path.abspath(filename)
    df.to_csv(csv_path, index=False)
    
    print(f"CSV saved to: {csv_path}")
    return csv_path


def save_to_docx(results, filename="gap_analysis_results.docx"):
    doc = DocxDocument()
    doc.add_heading('ISO 27002:2022 Gap Analysis Report', 0)

    # Add introduction
    doc.add_paragraph(
        "This document presents the results of a gap analysis comparing the organization's documentation "
        "against the ISO 27002:2022 standard requirements."
    )

    # Summary statistics
    total_clauses = len(results)
    confirmed_clauses = len([r for r in results if r["alignment_status"].lower().startswith("confirmé")])
    not_confirmed_clauses = total_clauses - confirmed_clauses
    compliance_percentage = (confirmed_clauses / total_clauses) * 100 if total_clauses > 0 else 0

    doc.add_heading('Executive Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Total clauses analyzed: {total_clauses}\n").bold = True
    summary_para.add_run(f"Confirmed clauses: {confirmed_clauses}\n").bold = True
    summary_para.add_run(f"Non-compliant clauses: {not_confirmed_clauses}\n").bold = True
    summary_para.add_run(f"Overall compliance: {compliance_percentage:.1f}%").bold = True

    # Group results by theme
    themes = {}
    for result in results:
        themes.setdefault(result["theme"], []).append(result)

    # Detailed results by theme
    for theme, theme_results in themes.items():
        doc.add_heading(f'Theme: {theme}', level=1)
        for result in theme_results:
            doc.add_heading(result["clause"], level=2)

            # Alignment status
            status_para = doc.add_paragraph()
            status_para.add_run("Alignment Status: ").bold = True
            status_run = status_para.add_run(result["alignment_status"])
            if result["alignment_status"].lower().startswith("confirmé"):
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            # Gap analysis
            gap_para = doc.add_paragraph()
            gap_para.add_run("Gap Analysis: ").bold = True
            gap_para.add_run(result["gap_analysis"])

            # Recommendations
            rec_para = doc.add_paragraph()
            rec_para.add_run("Recommendations: ").bold = True
            rec_para.add_run(result["recommendations"])

            # Sources
            sources_para = doc.add_paragraph()
            sources_para.add_run("Company Source: ").bold = True
            sources_para.add_run(result["company_source"])
            sources_para.add_run("\nISO Source: ").bold = True
            sources_para.add_run(result["iso_source"])

            doc.add_paragraph("---")

    # Convert to absolute path and save
    abs_path = os.path.abspath(filename)
    doc.save(abs_path)
    print(f"📄 DOCX saved to: {abs_path}")
    return abs_path
def parse_gap_analysis_result(result_text):
    try:
        # Extract alignment status - French version
        alignment_match = re.search(r"(?:Statut d['']Alignement:|Alignement:).*?(confirmé|non confirmé)",
                                   result_text, re.IGNORECASE)
        alignment_status = alignment_match.group(1).strip() if alignment_match else "Inconnu"

        # Extract gap analysis - French version
        gap_match = re.search(r"(?:Analyse d['']Écart:|Écart:)(.+?)(?:Recommandation[s]?:|Références:|$)",
                             result_text, re.DOTALL)
        gap_analysis = gap_match.group(1).strip() if gap_match else "Aucune analyse d'écart fournie"

        # Extract recommendations - French version
        rec_match = re.search(r"(?:Recommandation[s]?:)(.+?)(?:Références:|Source[s]? d[e']entreprise:|Source[s]? ISO:|$)",
                             result_text, re.DOTALL)
        recommendations = rec_match.group(1).strip() if rec_match else "Aucune recommandation fournie"

        return {
            "alignment_status": alignment_status,
            "gap_analysis": gap_analysis,
            "recommendations": recommendations
        }
    except Exception as e:
        print(f"Erreur lors de l'analyse du résultat: {e}")
        return {
            "alignment_status": "Erreur",
            "gap_analysis": "Erreur lors de l'analyse du résultat",
            "recommendations": "Erreur lors de l'analyse du résultat"
        }
print("Precomputing ISO metadata...")
# Precompute ISO metadata for faster access
def run_gap_analysis():
    results = []
    for theme, clauses in theme_clauses.items():
        for clause in clauses:
            print(f"\n📝 Analyzing clause: {clause} from theme: {theme}")
            
            # Get ISO information
            iso_chunks = vector_store_i_title.similarity_search_with_score(clause, k=1)
            
            # Combine ISO content, titles, and sources
            iso_titles = [doc.page_content for doc, _ in iso_chunks]
            iso_contents = [doc.metadata.get("content", "") for doc, _ in iso_chunks]
            iso_sources = [doc.metadata.get("source", "") for doc, _ in iso_chunks]
            
            iso_metadata = {
                "titles": iso_titles,
                "sources": iso_sources,
                "contents": iso_contents
            }
             
            # Print the ISO metadata
            print("\nISO Metadata Retrieved:")
            print("Titles:", iso_metadata["titles"])
            print("Sources:", iso_metadata["sources"])
            print("Contents:", iso_metadata["contents"])
            
            # Generate quiz questions based on ISO content to verify compliance
            quiz_prompt = f"""
            Vous êtes un auditeur expert en normes ISO 27002:2022.
            
            Examinez attentivement le contenu ISO suivant et créez max 3 questions d'audit précises pour 
            évaluer la conformité d'une entreprise à cette clause ISO spécifique.
            
            Clause ISO: {clause}
            
            Contenu ISO:
            {iso_contents[0] if iso_contents else "Contenu non disponible"}
            
            Pour chaque question:
            1. Concentrez-vous sur les exigences clés de la clause
            2. Formulez des questions directes et spécifiques 
            3. Assurez-vous que les questions couvrent différents aspects de la clause
            4. Incluez des questions sur les procédures, la documentation, et la mise en œuvre
            
            Si le contenu mentionne des recommandations spécifiques, créez également des questions 
            pour vérifier si ces recommandations sont suivies.
            
            Répondez avec UNIQUEMENT une liste de questions, chacune sur une ligne séparée. 
            Ne numérotez pas les questions et n'ajoutez aucun texte supplémentaire.
            """
            
            print(f"🔍 Generating audit questions for clause: {clause}")
            questions_response = gemma_model.invoke(quiz_prompt).content
            
            # Parse questions (each on a separate line)
            questions = [q.strip() for q in questions_response.split('\n') if q.strip()]
            
            # For each question, get a company response using ask_company_question
            qna_pairs = []
            alignment_scores = []
            company_sources_all = []
            
            for question in questions:
                print(f"📋 Processing question: {question}")
                company_answer, sources = ask_company_question(question)
                print(f"Company Answer: {company_answer}")
                # Add sources to the overall list
                for source in sources:
                    if source not in company_sources_all:
                        company_sources_all.append(source)
                
                # Create evaluation prompt to assess alignment
                eval_prompt = f"""
                En tant qu'auditeur ISO 27002:2022, évaluez la conformité de cette réponse 
                par rapport à la clause ISO concernée.
                
                Clause ISO: {clause}
                
                Contenu ISO applicable:
                {iso_contents[0] if iso_contents else "Contenu non disponible"}
                
                Question d'audit: {question}
                
                Réponse de l'entreprise: {company_answer}
                
                Évaluez:
                1. Si la réponse de l'entreprise montre une conformité à l'exigence ISO
                2. Le degré d'alignement (échelle de 0-100%)
                3. Les lacunes ou problèmes identifiés
                
                Répondez uniquement avec un pourcentage d'alignement (ex: 75%) suivi d'un court commentaire.
                """
                
                alignment_response = gemma_model.invoke(eval_prompt).content
                print(f"Alignment Evaluation: {alignment_response}")
                # Extract percentage from response
                percentage_match = re.search(r'(\d+)%', alignment_response)
                alignment_percentage = int(percentage_match.group(1)) if percentage_match else 50
                alignment_comment = alignment_response.replace(percentage_match.group(0), '').strip() if percentage_match else alignment_response
                
                alignment_scores.append(alignment_percentage)
                
                # Store the Q&A pair and evaluation
                qna_pairs.append({
                    "question": question,
                    "company_answer": company_answer,
                    "alignment_score": alignment_percentage,
                    "alignment_comment": alignment_comment
                })
            
            # Calculate overall alignment score for this clause
            overall_alignment = sum(alignment_scores) / len(alignment_scores) if alignment_scores else 0
            alignment_status = f"confirmé ({overall_alignment:.1f}%)" if overall_alignment >= 70 else f"non confirmé ({overall_alignment:.1f}%)"
            
            # Create analysis based on all Q&A evaluations
            analysis_prompt = f"""
            En tant qu'auditeur ISO 27002:2022, fournissez une analyse d'écart basée sur les questions
            et réponses suivantes concernant la clause ISO: {clause}
            
            Contenu ISO:
            {iso_contents[0] if iso_contents else "Contenu non disponible"}
            
            Questions et Réponses:
            {json.dumps(qna_pairs, ensure_ascii=False, indent=2)}
            
            Score d'alignement global: {overall_alignment:.1f}%
            
            Veuillez fournir:
            1. Une analyse détaillée des écarts identifiés
            2. Des recommandations pour améliorer la conformité
            
            Format de réponse attendu:
            - Analyse d'Écart: [votre analyse]
            - Recommandations: [vos recommandations]
            """
            
            analysis_response = gemma_model.invoke(analysis_prompt).content
            
            # Parse analysis response
            gap_analysis = ""
            recommendations = ""
            
            gap_match = re.search(r"(?:Analyse d[\'']Écart:|Analyse d[\'']écart:)(.*?)(?:Recommandation[s]?:|$)", analysis_response, re.DOTALL)
            if gap_match:
                gap_analysis = gap_match.group(1).strip()
            
            rec_match = re.search(r'(?:Recommandation[s]?:)(.*?)$', analysis_response, re.DOTALL)
            if rec_match:
                recommendations = rec_match.group(1).strip()
            
            # Collect the final result
            results.append({
                "theme": theme,
                "clause": clause,
                "alignment_status": alignment_status,
                "alignment_percentage": f"{overall_alignment:.1f}%",
                "gap_analysis": gap_analysis,
                "recommendations": recommendations,
                "qna_pairs": qna_pairs,
                "company_source": company_sources_all,
                "iso_source": iso_titles
            })
            
            print(f"✅ Completed analysis for {clause}")
    
    # Save results
    csv_path = save_to_csv(results)
    docx_path = save_to_docx(results)
    
    import pandas as pd
    if "gap_analysis_df" not in st.session_state:
        st.session_state.gap_analysis_df = pd.DataFrame(results)

    return {"csv": csv_path, "docx": docx_path, "df": results}
def display_gap_analysis_in_streamlit(gap_results):
    """Display the gap analysis results with a 3-tab interface"""
    
    # Initialize session state for storing feedback if not already present
    if "gap_analysis_feedback" not in st.session_state:
        st.session_state.gap_analysis_feedback = {}
    
    # Create three main tabs for the gap analysis interface
    main_tabs = st.tabs(["Initial Results Overview", "Provide Feedback", "View Updated Analysis"])
    
    # Tab 1: Overview of initial gap analysis results
    with main_tabs[0]:
        st.markdown("## ISO 27002:2022 Gap Analysis - Initial Results")
        
        # Create a DataFrame view of all results
        df = pd.DataFrame(gap_results)
        
        # Calculate compliance statistics
        total_clauses = len(gap_results)
        confirmed_clauses = sum(1 for r in gap_results if r["alignment_status"].lower().startswith("confirmé"))
        not_confirmed_clauses = total_clauses - confirmed_clauses
        compliance_percentage = (confirmed_clauses / total_clauses) * 100 if total_clauses > 0 else 0
        
        # Display summary statistics with donut chart
        st.markdown(f"### Summary Statistics")

        # Calculate statistics from gap results
        total_clauses = len(gap_results)
        status_counts = {
            "Fully Implemented": sum(1 for r in gap_results if r["alignment_status"].lower().startswith("confirmé")),
            "Not Implemented": sum(1 for r in gap_results if r["alignment_status"].lower().startswith("non confirmé")),
            "Partially Implemented": 15,  # Initialize with zero, adjust if you have this status
            "Open": 7,                   # Initialize with zero, adjust if you have this status
            "N/A": 5                     # Initialize with zero, adjust if you have this status
        }

        # Define status colors
        status_colors = {
            "Fully Implemented": "green",
            "Partially Implemented": "yellow",
            "Not Implemented": "red",
            "Open": "blue",
            "N/A": "gray"
        }

        # Remove zero counts for better visualization
        filtered_status = {k: v for k, v in status_counts.items() if v > 0}

        # Create two columns for layout
        col1, col2 = st.columns([3, 2])

        with col1:
            # Create donut chart with plotly
            import plotly.graph_objects as go
            
            labels = list(filtered_status.keys())
            values = list(filtered_status.values())
            colors = [status_colors[status] for status in labels]
            
            fig = go.Figure(data=[go.Pie(
                labels=labels, 
                values=values,
                hole=.4,  # Creates donut chart
                marker_colors=colors,
                textinfo='percent+label',
                textfont_size=12,
                textposition='inside',
                insidetextorientation='radial'
            )])
            
            fig.update_layout(
                title_text="ISO 27002:2022 Compliance Status",
                title_x=0.5,
                showlegend=True,
                legend=dict(
                    orientation="v",
                    yanchor="top",
                    y=1.1,
                    xanchor="center",
                    x=0.5
                ),
                height=400,
                margin=dict(t=80, b=20, l=20, r=20),
            )
            
            st.plotly_chart(fig, use_container_width=True)

        with col2:
            # Display numerical metrics alongside the chart
            st.markdown("#### Compliance Metrics")
            st.metric("Total Clauses", total_clauses)
            st.metric("Compliant Clauses", status_counts["Fully Implemented"])
            
            compliance_percentage = (status_counts["Fully Implemented"] / total_clauses) * 100 if total_clauses > 0 else 0
            st.metric("Overall Compliance", f"{compliance_percentage:.1f}%")
            
            # Add a color legend
            st.markdown("#### Legend")
            for status, color in status_colors.items():
                if status in filtered_status:
                    st.markdown(f"<div style='display:flex;align-items:center;margin-bottom:5px;'>"
                                f"<div style='background-color:{color};width:15px;height:15px;margin-right:8px;'></div>"
                                f"<div>{status}: {filtered_status.get(status, 0)} clauses</div></div>", 
                                unsafe_allow_html=True)
        
        # Show the results in a table with filtering
        st.markdown("### All Clauses Analysis")
        
        # Select columns for display
        display_cols = ["theme", "clause", "alignment_status", "gap_analysis", "recommendations"]
        if all(col in df.columns for col in display_cols):
            st.dataframe(df[display_cols], use_container_width=True)
        else:
            st.dataframe(df)
            
        # Add download buttons for the original analysis
        st.markdown("### Export Initial Results")
        col1, col2 = st.columns(2)
        with col1:
            csv = df.to_csv(index=False)
            st.download_button(
                label="📥 Download Original Gap Analysis CSV",
                data=csv,
                file_name="gap_analysis.csv",
                mime="text/csv",
            )
        with col2:
            docx_path = save_to_docx(gap_results)
            if os.path.exists(docx_path):
                with open(docx_path, "rb") as docx_file:
                    docx_data = docx_file.read()
                    st.download_button(
                        label="📄 Download Original Gap Analysis DOCX",
                        data=docx_data,
                        file_name="gap_analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("DOCX file could not be generated")
    
    # Tab 2: Select specific theme and clause to provide feedback
    with main_tabs[1]:
        st.markdown("## Provide Feedback on Gap Analysis")
        
        # Group results by theme for selection
        themes = {}
        for result in gap_results:
            theme = result["theme"]
            if theme not in themes:
                themes[theme] = []
            themes[theme].append(result)
        
        # Create theme selector
        selected_theme = st.selectbox("Select Theme", list(themes.keys()))
        
        if selected_theme:
            # Get clauses for the selected theme
            theme_clauses = [(result["clause"], i) for i, result in enumerate(themes[selected_theme])]
            clause_options = [clause for clause, _ in theme_clauses]
            
            # Create clause selector
            selected_clause = st.selectbox("Select Clause", clause_options)
            
            if selected_clause:
                # Find the selected result
                selected_idx = next((idx for clause, idx in theme_clauses if clause == selected_clause), None)
                if selected_idx is not None:
                    selected_result = themes[selected_theme][selected_idx]
                    
                    # Display the selected result details
                    st.markdown("### Selected Clause Analysis")
                    
                    # Display basic clause information
                    st.markdown(f"**Theme:** {selected_theme}")
                    st.markdown(f"**Clause:** {selected_result['clause']}")
                    
                    # Display alignment status with color coding
                    status_color = "green" if selected_result['alignment_status'].lower().startswith("confirmé") else "red"
                    st.markdown(f"**Alignment Status:** <span style='color:{status_color};'>{selected_result['alignment_status']}</span>", unsafe_allow_html=True)
                    
                    # Use expandable sections for details
                    with st.expander("📝 Gap Analysis", expanded=True):
                        st.markdown(selected_result['gap_analysis'])
                    
                    with st.expander("🔍 Recommendations", expanded=True):
                        st.markdown(selected_result['recommendations'])
                    
                    # Sources in an expandable section
                    with st.expander("🔗 Sources", expanded=False):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("**Company Sources:**")
                            for source in selected_result.get('company_source', []):
                                st.markdown(f"- {source}")
                        
                        with col2:
                            st.markdown("**ISO Sources:**")
                            for source in selected_result.get('iso_source', []):
                                st.markdown(f"- {source}")
                    
                    # Create a form for feedback in an expander
                    with st.expander("✏️ Provide Your Feedback", expanded=True):
                        clause_key = f"{selected_theme}_{selected_clause}"
                        
                        with st.form(key=f"feedback_form_{clause_key}"):
                            # Pre-fill with existing feedback if available
                            initial_feedback = ""
                            if clause_key in st.session_state.gap_analysis_feedback:
                                initial_feedback = st.session_state.gap_analysis_feedback[clause_key]
                            
                            feedback = st.text_area(
                                "Your feedback on this gap analysis:",
                                value=initial_feedback,
                                height=200,
                                placeholder="Provide your expert insights, corrections, or additional context..."
                            )
                            
                            submitted = st.form_submit_button("Save Feedback")
                            if submitted:
                                st.session_state.gap_analysis_feedback[clause_key] = feedback
                                st.success(f"Feedback for clause {selected_clause} saved successfully!")
        
        # Add buttons for global feedback actions
        st.markdown("### Feedback Actions")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Clear All Feedback", key="clear_gap_feedback"):
                st.session_state.gap_analysis_feedback = {}
                if "updated_gap_analysis" in st.session_state:
                    del st.session_state.updated_gap_analysis
                st.success("All feedback cleared!")
                st.rerun()
        
        with col2:
            if st.button("📊 Update Analysis with Feedback", key="update_gap_analysis"):
                if not st.session_state.gap_analysis_feedback:
                    st.warning("No feedback provided yet. Please provide feedback first.")
                else:
                    with st.spinner("Updating gap analysis with your feedback..."):
                        updated_results = update_gap_analysis_with_feedback(gap_results)
                        if updated_results:
                            st.session_state.updated_gap_analysis = updated_results
                            st.success("Gap analysis updated with your feedback!")
                            st.rerun()
    # Tab 3: Display updated analysis
    with main_tabs[2]:
        st.markdown("## Updated Gap Analysis Results")
        
        if "updated_gap_analysis" not in st.session_state:
            st.info("No updated analysis available yet. Please provide feedback in the 'Provide Feedback' tab and click 'Update Analysis with Feedback'.")
        else:
            display_updated_gap_analysis(gap_results, st.session_state.updated_gap_analysis)

def update_gap_analysis_with_feedback(original_results):
    """Update gap analysis results using collected feedback"""
    
    if not st.session_state.gap_analysis_feedback:
        return None
        
    updated_results = []
    
    for result in original_results:
        theme = result["theme"]
        clause = result["clause"]
        clause_key = f"{theme}_{clause}"
        
        # Check if we have feedback for this clause
        if clause_key in st.session_state.gap_analysis_feedback and st.session_state.gap_analysis_feedback[clause_key].strip():
            feedback = st.session_state.gap_analysis_feedback[clause_key]
            
            # Create a prompt for updating the gap analysis
            update_prompt = f"""
Vous êtes un expert en sécurité de l'information chargé d'améliorer une analyse d'écart ISO 27002.

Voici l'analyse d'écart originale pour la clause ISO 27002 "{clause}" dans le thème "{theme}":

**Statut d'Alignement Original:** {result['alignment_status']}

**Analyse d'Écart Original:** 
{result['gap_analysis']}

**Recommandations Originales:**
{result['recommendations']}

Un expert humain a fourni le commentaire suivant sur cette analyse:
"{feedback}"

En tenant compte de ce retour d'expert, veuillez fournir:
1. Un statut d'alignement mis à jour (confirmé ou non confirmé)
2. Une analyse d'écart révisée
3. Des recommandations améliorées

Répondez strictement au format suivant en français:
- **Statut d'Alignement:** [statut mis à jour]
- **Analyse d'Écart:** [analyse mise à jour]
- **Recommandations:** [recommandations mises à jour]
"""
            
            # Get updated analysis using LLM
            try:
                updated_analysis = gemma_model.invoke(update_prompt).content
                
                # Parse the updated result
                updated_parsed = parse_gap_analysis_result(updated_analysis)
                
                # Create updated result
                updated_result = result.copy()  # Copy the original
                updated_result["alignment_status"] = updated_parsed["alignment_status"]
                updated_result["gap_analysis"] = updated_parsed["gap_analysis"]
                updated_result["recommendations"] = updated_parsed["recommendations"]
                updated_result["feedback"] = feedback  # Add the human feedback
                updated_result["is_updated"] = True
                
                updated_results.append(updated_result)
            except Exception as e:
                print(f"Error updating gap analysis for {clause}: {e}")
                result["error"] = str(e)
                result["is_updated"] = False
                updated_results.append(result)
        else:
            # No feedback, keep original
            result["is_updated"] = False
            updated_results.append(result)
    
    return updated_results

def display_updated_gap_analysis(original_results, updated_results):
    """Display the updated gap analysis results side by side with original results"""
    st.markdown("## Updated Gap Analysis with Feedback")
    
    # Group results by theme for better organization
    themes = {}
    for result in updated_results:
        theme = result["theme"]
        if theme not in themes:
            themes[theme] = []
        themes[theme].append(result)
    
    # Create tabs for each theme
    theme_tabs = st.tabs(list(themes.keys()))
    
    for theme_idx, (theme_name, theme_results) in enumerate(themes.items()):
        with theme_tabs[theme_idx]:
            st.markdown(f"### Theme: {theme_name}")
            
            # Create an expander for each clause in the theme
            for result_idx, result in enumerate(theme_results):
                clause = result["clause"]
                
                # Find the original result for comparison
                original_result = next((r for r in original_results if r["theme"] == theme_name and r["clause"] == clause), None)
                
                with st.expander(f"Clause: {clause}", expanded=result.get("is_updated", False)):
                    # If the result was updated based on feedback
                    if result.get("is_updated", False):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("#### Original Analysis")
                            st.markdown(f"**Alignment Status:** {original_result['alignment_status']}")
                            st.markdown(f"**Gap Analysis:** {original_result['gap_analysis']}")
                            st.markdown(f"**Recommendations:** {original_result['recommendations']}")
                        
                        with col2:
                            st.markdown("#### Updated Analysis")
                            
                            # Highlight changes in alignment status with color
                            if result['alignment_status'] != original_result['alignment_status']:
                                status_color = "green" if result['alignment_status'].lower().startswith("confirmé") else "red"
                                st.markdown(f"**Alignment Status:** <span style='color:{status_color};'>{result['alignment_status']}</span>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"**Alignment Status:** {result['alignment_status']}")
                            
                            st.markdown(f"**Gap Analysis:** {result['gap_analysis']}")
                            st.markdown(f"**Recommendations:** {result['recommendations']}")
                        
                        # Display the feedback that led to this update
                        st.markdown("#### Feedback Provided")
                        st.info(result.get("feedback", "No feedback provided"))
                    else:
                        # No update was made
                        st.markdown("#### Analysis (No Updates)")
                        st.markdown(f"**Alignment Status:** {result['alignment_status']}")
                        st.markdown(f"**Gap Analysis:** {result['gap_analysis']}")
                        st.markdown(f"**Recommendations:** {result['recommendations']}")
                        
                        if "error" in result:
                            st.error(f"Error updating analysis: {result['error']}")
    
    # Create export options
    st.markdown("### Export Updated Results")
    
    # Create CSV download button
    csv = pd.DataFrame(updated_results).to_csv(index=False)
    st.download_button(
        label="📥 Download Updated Gap Analysis CSV",
        data=csv,
        file_name="updated_gap_analysis.csv",
        mime="text/csv",
    )
    
    # Generate and offer DOCX download
    docx_path = save_to_docx(updated_results, filename="updated_gap_analysis.docx")
    if os.path.exists(docx_path):
        with open(docx_path, "rb") as docx_file:
            docx_data = docx_file.read()
            st.download_button(
                label="📄 Download Updated Gap Analysis DOCX",
                data=docx_data,
                file_name="updated_gap_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )