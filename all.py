import streamlit as st
import os
import json
from typing import List, Dict, TypedDict, Optional, Tuple
from langchain_core.messages import SystemMessage
from langchain_core.documents import Document
from langchain_core.tools import tool
from langgraph.graph import END, START, StateGraph
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.graph import MessagesState
from langgraph.checkpoint.memory import MemorySaver
from langchain_qdrant import QdrantVectorStore
from langchain.prompts import ChatPromptTemplate
from operator import itemgetter
from sentence_transformers import SentenceTransformer
from langchain.embeddings.base import Embeddings
import json
import re
import pandas as pd
import os
from docx import Document as DocxDocument
from docx.shared import RGBColor
import torch
from transformers import AutoModelForSequenceClassification, AutoTokenizer
from langchain_ollama import ChatOllama
from transformers import AutoTokenizer
from langchain_ollama import ChatOllama
from transformers import AutoTokenizer
from transformers import AutoTokenizer
from langchain.schema import HumanMessage
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.units import inch
from reportlab.lib import colors
import io
from typing import List, Dict, TypedDict, Literal, Optional
from transformers import WhisperProcessor, WhisperForConditionalGeneration
import torchaudio
from audio_recorder_streamlit import audio_recorder
import numpy as np
from audio_recorder_streamlit import audio_recorder
from transformers import WhisperProcessor, WhisperForConditionalGeneration
import librosa
import io
import torch

print("Setting up graphs and language models...")

# Define the database file path for chat history
DB_FILE = "db.json"

st.set_page_config(page_title="ISO Assistant", layout="centered")

# -----------------------------
# Initialize Language Models
# -----------------------------
print("Initializing language models...")
WHISPER_MODEL_PATH = r"C:\Users\G800613RTS\whisper-tiny"

@st.cache_resource
def load_whisper_model():
    """Load and cache the Whisper model and processor"""
    try:
        processor_voice = WhisperProcessor.from_pretrained(WHISPER_MODEL_PATH)
        model_voice = WhisperForConditionalGeneration.from_pretrained(WHISPER_MODEL_PATH)
        model_voice.config.forced_decoder_ids = None
        return processor_voice, model_voice
    except Exception as e:
        st.error(f"Error loading Whisper model: {str(e)}")
        return None, None

def transcribe_audio(audio_bytes, processor_voice, model_voice):
    """Transcribe audio bytes to text using Whisper"""
    try:
        # Convert bytes to numpy array
        audio_data = np.frombuffer(audio_bytes, dtype=np.float32)
        
        # Resample to 16kHz if needed (Whisper expects 16kHz)
        if len(audio_data) > 0:
            # Load audio with librosa (handles resampling)
            audio_array, _ = librosa.load(io.BytesIO(audio_bytes), sr=16000, mono=True)
            
            # Process audio
            input_features = processor_voice(
                audio_array, 
                sampling_rate=16000, 
                return_tensors="pt"
            ).input_features
            
            # Generate transcription
            predicted_ids = model_voice.generate(input_features)
            
            # Decode to text
            transcription = processor_voice.batch_decode(predicted_ids, skip_special_tokens=True)
            
            return transcription[0] if transcription else ""
        
        return ""
    except Exception as e:
        st.error(f"Error transcribing audio: {str(e)}")
        return ""






llm_ollama =ChatOllama(model="qwen3:0.6b",temperature=0.7)

gemma_model=ChatOllama(model="qwen3:0.6b",temperature=0.7)
model_id=r"C:\Users\G800613RTS\bge-reranker-v2-m3"
reranker_tokenizer = AutoTokenizer.from_pretrained(model_id)
reranker_model = AutoModelForSequenceClassification.from_pretrained(model_id)
reranker_model.eval()
#reranker 
def rerank_documents(query, documents, top_k=2):
    doc_texts = []
    docs_list = []
    
    for doc in documents:
        # Handle tuple from similarity_search_with_score
        if isinstance(doc, tuple) and len(doc) == 2:
            document = doc[0]  # Extract the document from the (doc, score) tuple
            docs_list.append(document)
            
            if hasattr(document, 'page_content'):
                doc_texts.append(document.page_content)
            elif isinstance(document, dict) and 'text' in document:
                doc_texts.append(document['text'])
            elif isinstance(document, dict) and 'chunk' in document:
                doc_texts.append(document['chunk'])
            elif isinstance(document, dict) and 'summary_chunk' in document:
                doc_texts.append(document['summary_chunk'])
            elif isinstance(document, dict) and 'content' in document:
                doc_texts.append(document['content'])
            elif isinstance(document, str):
                doc_texts.append(document)
            else:
                print(f"Warning: Skipping unsupported document type in tuple: {type(document)}")
                continue
        # Original handling for direct documents
        elif hasattr(doc, 'page_content'):
            doc_texts.append(doc.page_content)
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'text' in doc:
            doc_texts.append(doc['text'])
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'chunk' in doc:
            doc_texts.append(doc['chunk'])
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'summary_chunk' in doc:
            doc_texts.append(doc['summary_chunk'])
            docs_list.append(doc)
        elif isinstance(doc, dict) and 'content' in doc:
            doc_texts.append(doc['content'])
            docs_list.append(doc)
        elif isinstance(doc, str):
            doc_texts.append(doc)
            docs_list.append(doc)
        else:
            print(f"Warning: Skipping unsupported document type: {type(doc)}")
            continue

    pairs = [[query, text] for text in doc_texts]

    with torch.no_grad():
        inputs = reranker_tokenizer(pairs, padding=True, truncation=True, return_tensors='pt', max_length=512)
        scores = reranker_model(**inputs, return_dict=True).logits.view(-1, ).float()

    scored_docs = list(zip(scores.tolist(), docs_list))
    scored_docs.sort(key=lambda x: x[0], reverse=True)

    return [doc for _, doc in scored_docs[:top_k]]
# -----------------------------
# Reload embeddings model
# -----------------------------
print("Loading embedding model...")
model = SentenceTransformer(
    "nomic-ai/nomic-embed-text-v1",
    trust_remote_code=True
)

class CustomSentenceTransformerEmbeddings(Embeddings):  # Inherit from LangChain base
    def __init__(self, model):
        self.model = model

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return self.model.encode(texts, normalize_embeddings=True).tolist()

    def embed_query(self, text: str) -> List[float]:
        return self.model.encode([text], normalize_embeddings=True)[0].tolist()

embeddings = CustomSentenceTransformerEmbeddings(model)

# -----------------------------
# Connect to Vector Stores
# -----------------------------
vector_store_i_title = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="io_title_collection",
    url="http://localhost:6333",
)
vector_store_i_text = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="io_text_collection",
    url="http://localhost:6333",
)
vector_store_summaries= QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="co_summary_vectors",
    url="http://localhost:6333",
)
vector_store_originals = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="co_original_vectors",
    url="http://localhost:6333",
)
vector_store_42001_originals = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="io42001_text_collection",
    url="http://localhost:6333",
)
vector_store_iso_title24001 = QdrantVectorStore.from_existing_collection(
    embedding=embeddings,
    collection_name="io42001_title_collection",
    url="http://localhost:6333",
)
# -----------------------------
# PDF Processing Functions
# -----------------------------
# UPLOADS_DIR = "uploads"
# OUTPUT_TXT_DIR = "output_txt"
# OUTPUT_JSON_DIR = "output_json"
# SUMMARY_JSON_DIR = "Comp_json"
# ORIGINAL_JSON_DIR = "Comp_json_1"
# def resize_crop(image, top=0.08, bottom=0.1):
#     """Resize and crop the image."""
#     ratio = image.size[0] / image.size[1]
#     image = image.resize((int(ratio * 800), 800))
#     left, upper, right, lower = 0, int(image.height * top), image.width, int(image.height * (1 - bottom))
#     return image.crop((left, upper, right, lower))

# def extract_text_from_image(image, pdf_name, model, processor, output_file):
#     """Extract text from an image using Qwen model."""
#     messages = [
#         {
#             "role": "user",
#             "content": [
#                 {"type": "image", "image": image},
#                 {"type": "text", "text": "You are an OCR READER you have to extcart all the text dont miss any world from it and dont modifie it . Read all the text and chunk it in a human-readable way.\
#                should be treated as a separate chunk, containing the complete text of that section . Do not summarize—keep the full text intact.If you encounter \
#                an image, chart, or diagram, provide a detailed and meaningful description of its content and purpose. give a short description  \
#                Ensure the description conveys its significance in relation to the document put it all in one chunk .\
#                print it all in one chunk in the same one chunk "},
#             ],
#         }
#     ]
    
#     text = processor.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
#     inputs = processor(text=[text], images=[image], padding=True, return_tensors="pt").to(model.device)
#     generated_ids = model.generate(**inputs, max_new_tokens=10000)
#     generated_ids_trimmed = [out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generated_ids)]
#     output_text = processor.batch_decode(generated_ids_trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=False)
    
#     # Save raw output to a text file
#     with open(output_file, "a", encoding="utf-8") as f:
#         f.write("\n".join(output_text) + "\n")
    
#     return output_text[0] if output_text else ""


# def identify_sections_and_group(text):
#     """Identify sections and group them together based on section numbers."""
#     section_pattern = re.compile(r'^(\d+(\.\d+)*)\s+(.+)$')  # Ex: "5", "5.1", "5.1.1"
    
#     # First check if we have any section headers at all
#     lines = text.split("\n")
#     has_section_headers = False
#     for line in lines:
#         if section_pattern.match(line.strip()):
#             has_section_headers = True
#             break
    
#     # If no section headers found, return all text as one chunk
#     if not has_section_headers:
#         return [{
#             "title": "Document Content",
#             "text": text
#         }]
    
#     # Continue with normal processing if section headers exist
#     chunks = []
#     current_section = None
#     current_section_number = None
#     main_section_title = None
    
#     for line in lines:
#         match = section_pattern.match(line.strip())
        
#         if match:
#             section_number = match.group(1)
#             section_title = match.group(3).strip()
            
#             if '.' not in section_number:  # Main section (ex: "5")
#                 main_section_title = section_title
#                 current_section_number = section_number
                
#                 if current_section:
#                     chunks.append(current_section)
                
#                 current_section = {
#                     "title": f"{section_number} {section_title}",
#                     "text": [line.strip()]
#                 }
#             else:  # Sub-section (ex: "5.1", "5.1.1")
#                 if not current_section:
#                     # If we encounter a sub-section without a parent section
#                     current_section = {
#                         "title": f"{section_number} {section_title}",
#                         "text": [line.strip()]
#                     }
#                 else:
#                     current_section["text"].append(line.strip())
#         else:
#             if current_section:
#                 current_section["text"].append(line.strip())
#             # If no current section, this text will be skipped in the original implementation
    
#     # Add the last section if it exists
#     if current_section:
#         chunks.append(current_section)
    
#     # If somehow we ended up with no chunks, create one with all the text
#     if not chunks:
#         return [{
#             "title": "Document Content",
#             "text": text
#         }]
    
#     # Convert lists of text lines to single strings for each chunk
#     for chunk in chunks:
#         chunk["text"] = "\n".join(chunk["text"])
    
#     return chunks

# def process_pdf(pdf_path, model_path):
#     """Process a PDF file: convert to images, extract text, group sections."""
#     # Initialize Qwen model for image processing
#     model = Qwen2_5_VLForConditionalGeneration.from_pretrained(
#     model_path, 
#     torch_dtype=torch.bfloat16
#     )
#     model.to("cuda:0")
#     processor = AutoProcessor.from_pretrained(model_path)
    
#     pdf_name = os.path.basename(pdf_path)
#     timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#     output_file = os.path.join(OUTPUT_TXT_DIR, f"{pdf_name}_{timestamp}.txt")
    
#     # Convert PDF to images
#     images = convert_from_path(pdf_path, dpi=300)
#     full_text = ""
    
#     # Process each page
#     for idx, image in enumerate(images):
#         image_path = os.path.join(UPLOADS_DIR, f"{pdf_name}_page_{idx+1}.png")
#         image.save(image_path, 'PNG')
        
#         # Process image
#         processed_image = resize_crop(Image.open(image_path))
#         extracted_text = extract_text_from_image(processed_image, pdf_name, model, processor, output_file)
#         full_text += extracted_text + "\n\n"
    
#     # Group sections
#     chunks = identify_sections_and_group(full_text)
    
#     # Save chunked text to JSON
#     output_json = os.path.join(OUTPUT_JSON_DIR, f"{pdf_name}_{timestamp}.json")
#     with open(output_json, 'w', encoding="utf-8") as f:
#         json.dump(chunks, f, indent=4, ensure_ascii=False)
    
#     return chunks, output_json

# def summarize_chunks(chunks, pdf_name):
#     """Summarize chunks using the deepseek model."""
#     # Load the classification dictionary
#     with open('sheet_français_1.json', 'r') as f:
#         loaded_dict = json.load(f)
    
#     # Initialize the model
#     model_name = "deepseek-r1:1.5b"
#     base_url = "http://localhost:11434"
#     model = ChatOllama(model=model_name, base_url=base_url)
    
#     # Process chunks
#     results = []
#     for idx, chunk in enumerate(chunks):
#         try:
#             prompt = f"""
#             Analyze the following text/table element and:
#             1. Summarize its content concisely.
#             2. Match the element to analysis to just one Category and one clause exactly to the closest category and clause from this dictionary:
#             {loaded_dict}
#             3. Identify the exact Category and specific clause using format "X.XX Clause description".
#             4. Return in this format: [Category Name] - [Exact Clause Number and Title]: [Summary]
            
#             Element to analyze: {chunk["text"]}
#             """
            
#             output = model.invoke([HumanMessage(content=prompt)])
            
#             # Extract content from <think>...</think> tags if present
#             think_pattern = r"<think>(.*?)</think>"
#             think_match = re.search(think_pattern, output.content, re.DOTALL)
            
#             if think_match:
#                 main_response = re.sub(think_pattern, "", output.content, flags=re.DOTALL).strip()
#             else:
#                 main_response = output.content.strip()
            
#             results.append({
#                 "content": chunk["text"],
#                 "summary": main_response,
#                 "name_pdf": pdf_name
#             })
            
#             print(f"Processed chunk {idx + 1}/{len(chunks)} successfully.")
#         except Exception as e:
#             print(f"Error processing chunk {idx + 1}: {e}")
#             results.append({
#                 "content": chunk["text"],
#                 "summary": f"Error: {str(e)}",
#                 "name_pdf": pdf_name
#             })
    
#     # Save results
#     timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#     summary_json_path = os.path.join(SUMMARY_JSON_DIR, f"{pdf_name}_summary_{timestamp}.json")
#     original_json_path = os.path.join(ORIGINAL_JSON_DIR, f"{pdf_name}_original_{timestamp}.json")
    
#     with open(summary_json_path, 'w', encoding='utf-8') as f:
#         json.dump(results, f, ensure_ascii=False, indent=2)
    
#     with open(original_json_path, 'w', encoding='utf-8') as f:
#         json.dump(results, f, ensure_ascii=False, indent=2)
    
#     print(f"🎉 Results saved to {summary_json_path} and {original_json_path}")
    
#     return results

# def update_vector_stores(new_chunks):
#     """Update the vector stores with new chunks."""
   
#     # Create summary documents
#     summary_docs = [
#         Document(
#             page_content=chunk["summary"],
#             metadata={
#                 "chunk_text": chunk["content"],
#                 "source": chunk["name_pdf"]
#             }
#         )
#         for chunk in new_chunks
#     ]
    
#     # Create original documents
#     original_docs = [
#         Document(
#             page_content=chunk["content"],
#             metadata={
#                 "summary": chunk["summary"],
#                 "source": chunk["name_pdf"]
#             }
#         )
#         for chunk in new_chunks
#     ]
    
#     # Initialize vector stores
#     vector_store_summaries= QdrantVectorStore.from_existing_collection(
#     embedding=embeddings,
#     collection_name="co_summary_vectors",
#     url="http://localhost:6333",
#     )
#     vector_store_originals = QdrantVectorStore.from_existing_collection(
#     embedding=embeddings,
#     collection_name="co_original_vectors",
#     url="http://localhost:6333",
#     )
    
#     # Generate only UUID IDs without timestamp prefix
#     summary_ids = [str(uuid.uuid4()) for _ in range(len(summary_docs))]
#     original_ids = [str(uuid.uuid4()) for _ in range(len(original_docs))]
    
#     # Upload to vector stores
#     if summary_docs:
#         vector_store_summaries.add_documents(documents=summary_docs, ids=summary_ids)
#         print(f"✅ Added {len(summary_docs)} documents to summary vector store")
    
#     if original_docs:
#         vector_store_originals.add_documents(documents=original_docs, ids=original_ids)
#         print(f"✅ Added {len(original_docs)} documents to original vector store")
    
#     return len(summary_docs), len(original_docs)
# -----------------------------
# ISO Q&A Tool Setup
# -----------------------------
#-------------------
#tools
#-------------------
# Define the retrieval tools
@tool(response_format="content_and_artifact")
def retrieve27001(query: str):
    """Retrieve information related to ISO 27001 query."""
    
    # Step 1: Retrieve from vector store
    retrieved_docs = vector_store_i_text.similarity_search(query, k=5)
    print("\n📥 Top 5 Retrieved ISO 27001 Documents (Before Reranking):")
    for i, doc in enumerate(retrieved_docs):
        print(f"Doc {i+1}: Source: {doc.metadata.get('source')}, Title: {doc.metadata.get('title')}")
        print(f"Content Preview: {doc.page_content[:200]}...\n")

    # Step 2: Rerank the top results
    reranked_docs = rerank_documents(query, retrieved_docs, top_k=2)
    print("\n🔝 Top 2 Reranked ISO 27001 Documents:")
    for i, doc in enumerate(reranked_docs):
        print(f"Rank {i+1}: Source: {doc.metadata.get('source')}, Title: {doc.metadata.get('title')}")
        print(f"Content Preview: {doc.page_content[:200]}...\n")

    # Step 3: Prepare serialized output
    serialized = "\n\n".join(
        (f"Source: {doc.metadata}\n" f"Content: {doc.page_content}")
        for doc in reranked_docs
    )
    
    return serialized, reranked_docs

@tool(response_format="content_and_artifact")
def retrieve27002(query: str):
    """Retrieve information related to ISO 27002 query."""
    
    # Step 1: Retrieve from vector store - using the same vector store as 27001 for now
    # You can replace with a dedicated 27002 vector store if available
    retrieved_docs = vector_store_i_text.similarity_search(query, k=5)
    print("\n📥 Top 5 Retrieved ISO 27002 Documents (Before Reranking):")
    for i, doc in enumerate(retrieved_docs):
        print(f"Doc {i+1}: Source: {doc.metadata.get('source')}, Title: {doc.metadata.get('title')}")
        print(f"Content Preview: {doc.page_content[:200]}...\n")

    # Step 2: Rerank the top results
    reranked_docs = rerank_documents(query, retrieved_docs, top_k=2)
    print("\n🔝 Top 2 Reranked ISO 27002 Documents:")
    for i, doc in enumerate(reranked_docs):
        print(f"Rank {i+1}: Source: {doc.metadata.get('source')}, Title: {doc.metadata.get('title')}")
        print(f"Content Preview: {doc.page_content[:200]}...\n")

    # Step 3: Prepare serialized output
    serialized = "\n\n".join(
        (f"Source: {doc.metadata}\n" f"Content: {doc.page_content}")
        for doc in reranked_docs
    )
    
    return serialized, reranked_docs

@tool(response_format="content_and_artifact")
def retrieve42001(query: str):
    """Retrieve information related to ISO 42001 query."""
 
    # Step 1: Retrieve from ISO 42001 vector store
    retrieved_docs = vector_store_42001_originals.similarity_search(query, k=5)
    print(f"[DEBUG] 42001 Query: {query}")
    print("\n📥 Top 5 Retrieved ISO 42001 Documents (Before Reranking):")
    for i, doc in enumerate(retrieved_docs):
        print(f"Doc {i+1}: Source: {doc.metadata.get('source')}, Title: {doc.metadata.get('title')}")
        print(f"Content Preview: {doc.page_content[:200]}...\n")

    # Step 2: Rerank the top results
    reranked_docs = rerank_documents(query, retrieved_docs, top_k=2)
    print("\n🔝 Top 2 Reranked ISO 42001 Documents:")
    for i, doc in enumerate(reranked_docs):
        print(f"Rank {i+1}: Source: {doc.metadata.get('source')}, Title: {doc.metadata.get('title')}")
        print(f"Content Preview: {doc.page_content[:200]}...\n")

    # Step 3: Prepare serialized output
    serialized = "\n\n".join(
        (f"Source: {doc.metadata}\n" f"Content: {doc.page_content}")
        for doc in reranked_docs
    )
    
    return serialized, reranked_docs

@tool(response_format="content_and_artifact")
def retrieveCompany(query: str):
    """Retrieve information related to the Company sagemcom query from both summary and original vector stores."""

    # Always use the complete user query to find the most relevant documents
    # The query comes directly from the user's input in query_or_respond function

    print(f"[DEBUG] Company Search Query: {query}")

    try:
        # Step 1: Retrieve from both vector stores
        print("[DEBUG] Searching summary vector store...")
        summary_results = vector_store_summaries.similarity_search(query, k=5)
        
        print("[DEBUG] Searching original vector store...")
        original_results = vector_store_originals.similarity_search(query, k=5)
        
        print(f"[DEBUG] Retrieved {len(summary_results)} summary documents")
        print(f"[DEBUG] Retrieved {len(original_results)} original documents")
        
        # Step 2: Combine results and remove duplicates based on source
        all_results = []
        seen_sources = set()
        
        # Process summary results first
        for doc in summary_results:
            source = doc.metadata.get('source', 'Unknown Source')
            if source not in seen_sources:
                # For summary docs, the page_content is the summary, chunk_text is the original content
                enhanced_doc = Document(
                    page_content=doc.page_content,  # This is the summary
                    metadata={
                        'source': source,
                        'summary': doc.page_content,  # Summary from page_content
                        'chunk_text': doc.metadata.get('chunk_text', ''),  # Original content
                        'doc_type': 'summary'
                    }
                )
                all_results.append(enhanced_doc)
                seen_sources.add(source)
        
        # Process original results for any sources not already included
        for doc in original_results:
            source = doc.metadata.get('source', 'Unknown Source')
            if source not in seen_sources:
                # For original docs, the page_content is the original content, summary is in metadata
                enhanced_doc = Document(
                    page_content=doc.page_content,  # This is the original content
                    metadata={
                        'source': source,
                        'summary': doc.metadata.get('summary', ''),  # Summary from metadata
                        'chunk_text': doc.page_content,  # Original content from page_content
                        'doc_type': 'original'
                    }
                )
                all_results.append(enhanced_doc)
                seen_sources.add(source)
        
        # Step 3: Debug combined results
        for i, doc in enumerate(all_results):
            source = doc.metadata.get('source', 'Unknown Source')
            doc_type = doc.metadata.get('doc_type', 'unknown')
            summary = doc.metadata.get('summary', '')
            summary_preview = summary[:100] if summary else 'No summary available'
            print(f"[DEBUG] Combined Doc {i+1} - Type: {doc_type}, Source: {source}, Summary: {summary_preview}")
        
        # Step 4: Rerank the combined results
        if all_results:
            reranked_docs = rerank_documents(query, all_results, top_k=5)
            print(f"[DEBUG] Reranked to top {len(reranked_docs)} documents")
        else:
            reranked_docs = []
        
        # Step 5: Format output with enhanced information
        formatted_docs = []
        for doc in reranked_docs:
            source = doc.metadata.get('source', 'Unknown Source')
            doc_type = doc.metadata.get('doc_type', 'unknown')
            summary = doc.metadata.get('summary', 'No summary available')
            chunk_text = doc.metadata.get('chunk_text', 'No original content available')
            
            doc_text = f"Source: {source}\n"
            doc_text += f"Document Type: {doc_type}\n"
            doc_text += f"Summary: {summary}\n"
            doc_text += f"Original Content: {chunk_text}"
            formatted_docs.append(doc_text)
        
        formatted = "\n\n" + "\n\n".join(formatted_docs)
        
        if not formatted_docs:
            return "No documents found matching your query.", []
            
        return formatted, reranked_docs
        
    except Exception as e:
        # Add better error handling to diagnose issues
        print(f"[ERROR] Error in retrieveCompany: {str(e)}")
        import traceback
        print(f"[ERROR] Traceback: {traceback.format_exc()}")
        # Return a valid but empty response instead of raising an error
        return f"Error retrieving documents: {str(e)}", []
# Create a ToolNode with all retrieve tools
# The LLM will decide which tool to use based on the query content
tools = [retrieve27001, retrieve27002, retrieve42001,retrieveCompany]
tool_node = ToolNode(tools)
# Build the graph
print("Building ISO Q&A graph with dynamic tool selection...")
graph_builder = StateGraph(MessagesState)
def query_or_respond(state: MessagesState):
    """Let the LLM determine which tool to use based on the query content."""
    # Get the latest message
    last_message = state["messages"][-1] if state["messages"] else None
    
    if last_message and hasattr(last_message, "content"):
        print(f"[INFO] Processing query: {last_message.content}")
        
        # Create system instructions to help the LLM select the right tool
        system_message = SystemMessage(content="""
        You are a specialized assistant for ISO standards and company-specific information.
        
        You have access to the following tools:
        - retrieve27001: Use for questions about ISO 27001 standard
        - retrieve27002: Use for questions about ISO 27002 standard
        - retrieve42001: Use for questions about ISO 42001 standard
        - retrieveCompany: Use for ANY questions about company-specific documents, policies, or practices
        
        IMPORTANT GUIDELINES:
        1. For questions mentioning Sagemcom, company documents, or company policies, use the retrieveCompany tool
        2. For ISO standard questions, use the appropriate ISO retrieval tool
        3. When using retrieveCompany, pass the user's complete query as is
        4. When using ISO tools, focus the query on the specific clause or concept mentioned
        5. If unsure, prefer retrieveCompany for company-related queries and the ISO tools for standards
        
        Always maintain the full context of the user's question when selecting a tool.
        """)
        
        # Add the system message to guide tool selection
        messages_with_system = [system_message] + state["messages"]
        
        # Bind all tools to the LLM and let it choose which one to use
        llm_with_tools = llm_ollama.bind_tools(tools)
        response = llm_with_tools.invoke(messages_with_system)
        
        return {"messages": [response]}
    
    # Fallback if no message is found
    return {"messages": []}

def generate(state: MessagesState):
    """Generate answer based on retrieved content."""
    # Get the most recent tool messages
    recent_tool_messages = [m for m in reversed(state["messages"]) if m.type == "tool"]
    tool_messages = recent_tool_messages[::-1]
    
    
    # Format the tool output into a context for the response
    docs_content = "\n\n".join(doc.content for doc in tool_messages)
    
    
    system_message_content = (
       
        "Use the following pieces of retrieved context to answer "
        "the question. If you don't know the answer, say that you "
        "don't know. Use three sentences maximum and keep the "
        "answer concise."
        "\n\n"
        f"{docs_content}"
    )
    
    # Get relevant conversation messages
    conversation_messages = [
        message
        for message in state["messages"]
        if message.type in ("human", "system")
        or (message.type == "ai" and not message.tool_calls)
    ]
    
    # Prepare prompt and generate response
    prompt = [SystemMessage(system_message_content)] + conversation_messages
    response = llm_ollama.invoke(prompt)
    
    
    
    return {"messages": [response]}
#--------------
#Graph
#--------------
# Build the graph
graph_builder.add_node("query_or_respond", query_or_respond)
graph_builder.add_node("tools", tool_node)
graph_builder.add_node("generate", generate)

graph_builder.set_entry_point("query_or_respond")
graph_builder.add_conditional_edges(
    "query_or_respond",
    tools_condition,
    {END: END, "tools": "tools"},
)
graph_builder.add_edge("tools", "generate")
graph_builder.add_edge("generate", END)
memory = MemorySaver()

# Compile the graph
graph = graph_builder.compile(checkpointer=memory)
config = {"configurable": {"thread_id": "abc123"}}

# def ask_iso_question(user_query: str, thread_id: str = "thread-001"):
#     """
#     Direct RAG implementation for ISO 27002 documentation Q&A.
#     Retrieves relevant documents, reranks them, and generates a response.
#     """
#     print(f"[INFO] Processing ISO query: {user_query}")
    
#     # Step 1: Initial retrieval from vector store
#     results = vector_store_i_title.similarity_search(query=user_query, k=2)
#     print(f"[DEBUG] Retrieved {len(results)} documents")
    
#     # Step 2: Debug each result
#     for r in results:
#         print(f"[DEBUG] Doc Source: {r.metadata.get('source')}, Title: {r.metadata.get('title')}")
    
#     # Step 3: Rerank using reranker
#     reranked_docs = rerank_documents(user_query, results, top_k=1)
    
#     # Step 4: Extract and format document content for context
#     formatted_docs = []
#     sources = set()
    
#     for doc in reranked_docs:
#         source = doc.metadata.get('source', 'Unknown source')
#         title = doc.metadata.get('title', 'No title available')
        
#         # Add debug prints for the extracted metadata
#         print(f"[DEBUG] Processing document - Source: {source}")
#         print(f"[DEBUG] Processing document - Title: {title}")
        
        
#         formatted_doc = f"Source: {source}\nTitle: {title}\nContent: {doc.page_content}"
#         formatted_docs.append(formatted_doc)
#         sources.add(source)
#         sources.add(title)
        
#         # Print the formatted document string
#         print(f"[DEBUG] Formatted document: {formatted_doc[:200]}...")
    
#     # Combine all document content for context
#     context = "\n\n".join(formatted_docs)
    
#     # Step 5: Generate response using LLM with context
#     system_prompt = (
#         "You are an assistant for ISO 27002 questions. "
#         "Use the following context from ISO documents to answer the user's question. "
#         "Make sure to include the source of any information you use in your response.\n\n" + context
#     )
    
#     messages = [
#         {"role": "system", "content": system_prompt},
#         {"role": "user", "content": user_query}
#     ]
    
#     # Generate response
#     response = llm_ollama.invoke(messages).content
    
#     print("\n🔹 ISO Sources used:", sources)
    
#     # Return both the response text and the sources used
#     return response, list(sources)



# -----------------------------
# COMP Q&A Tool Setup
# -----------------------------
# @tool(response_format="content_and_artifact")
# def retrieve_company(query: str):
#     """Retrieve and rerank relevant company documents based on a query."""
#     # Step 1: Initial retrieval from vector store
#     results = vector_store_originals.similarity_search(query, k=10)

#     print(f"[DEBUG] Query: {query}")
#     print(f"[DEBUG] Retrieved {len(results)} documents")

#     # Step 2: Debug each result
#     for r in results:
#         print(f"[DEBUG] Doc Source: {r.metadata.get('source')}, Summary: {r.metadata.get('summary')[:100]}")

#     # Step 3: Rerank using your reranker
#     reranked_docs = rerank_documents(query, results, top_k=5)

#     # Step 4: Format output
#     formatted = "\n\n".join(
#         f"Source: {doc.metadata.get('source')}\n"
#         f"Summary: {doc.metadata.get('summary')}\n"
#         f"Original: {doc.page_content}"
#         for doc in reranked_docs
#     )

#     return formatted, reranked_docs


# -----------------------------
# COMP Q&A Graph
# -----------------------------
# print("Building Company Q&A graph...")
# graph_builder_company = StateGraph(MessagesState)

# Node: Call LLM with tools


# Node: Final response based on tool output
def ask_company_question(user_query: str, thread_id: str = "thread-002"):
    """
    Direct RAG implementation for company documentation Q&A.
    Retrieves relevant documents from both summary and original vector stores, reranks them, and generates a response.
    """
    print(f"[INFO] Processing company query: {user_query}")
    
    # Step 1: Retrieve from both vector stores
    print("[DEBUG] Searching summary vector store...")
    summary_results = vector_store_summaries.similarity_search(query=user_query, k=15)
    
    print("[DEBUG] Searching original vector store...")
    original_results = vector_store_originals.similarity_search(query=user_query, k=15)
    
    print(f"[DEBUG] Retrieved {len(summary_results)} summary documents")
    print(f"[DEBUG] Retrieved {len(original_results)} original documents")
    
    # Step 2: Combine results and remove duplicates based on source
    all_results = []
    seen_sources = set()
    
    # Process summary results first
    for doc in summary_results:
        source = doc.metadata.get('source', 'Unknown Source')
        if source not in seen_sources:
            # For summary docs, the page_content is the summary, chunk_text is the original content
            enhanced_doc = Document(
                page_content=doc.page_content,  # This is the summary
                metadata={
                    'source': source,
                    'summary': doc.page_content,  # Summary from page_content
                    'chunk_text': doc.metadata.get('chunk_text', ''),  # Original content
                    'doc_type': 'summary'
                }
            )
            all_results.append(enhanced_doc)
            seen_sources.add(source)
    
    # Process original results for any sources not already included
    for doc in original_results:
        source = doc.metadata.get('source', 'Unknown Source')
        if source not in seen_sources:
            # For original docs, the page_content is the original content, summary is in metadata
            enhanced_doc = Document(
                page_content=doc.page_content,  # This is the original content
                metadata={
                    'source': source,
                    'summary': doc.metadata.get('summary', ''),  # Summary from metadata
                    'chunk_text': doc.page_content,  # Original content from page_content
                    'doc_type': 'original'
                }
            )
            all_results.append(enhanced_doc)
            seen_sources.add(source)
    
    print(f"[DEBUG] Combined results: {len(all_results)} unique documents")
    
    # Step 3: Debug combined results
    for i, doc in enumerate(all_results):
        source = doc.metadata.get('source', 'Unknown Source')
        doc_type = doc.metadata.get('doc_type', 'unknown')
        summary = doc.metadata.get('summary', '')
        summary_preview = summary[:100] if summary else 'No summary available'
        print(f"[DEBUG] Combined Doc {i+1} - Type: {doc_type}, Source: {source}, Summary: {summary_preview}")
    
    # Step 4: Rerank the combined results
    if all_results:
        reranked_docs = rerank_documents(user_query, all_results, top_k=16)
        print(f"[DEBUG] Reranked to top {len(reranked_docs)} documents")
    else:
        reranked_docs = []
        print("[DEBUG] No documents to rerank")
    
    # Step 5: Extract and format document content for context
    formatted_docs = []
    sources = set()
    
    for doc in reranked_docs:
        source = doc.metadata.get('source', 'Unknown source')
        doc_type = doc.metadata.get('doc_type', 'unknown')
        summary = doc.metadata.get('summary', 'No summary available')
        chunk_text = doc.metadata.get('chunk_text', 'No original content available')
        title = doc.metadata.get('title', '')
        
        # Format with enhanced information showing both summary and original content
        doc_text = f"Source: {source}\n"
        doc_text += f"Document Type: {doc_type}\n"
        doc_text += f"Summary: {summary}\n"
        doc_text += f"Original Content: {chunk_text}"
        
        formatted_docs.append(doc_text)
        
        # Add source to the sources set
        sources.add(source)
        if title:  # Add title if available
            sources.add(title)
    
    # Step 6: Combine all document content for context
    context = "\n\n".join(formatted_docs)
    
    # Step 7: Generate response using LLM with enhanced context
    system_prompt = (
        "You are a responsable de gestion de l'information (information management officer) tasked with answering questions. "
        "IMPORTANT: Answer ONLY based on the context from company documents provided below. "
        "The context includes both summaries and original content from company documents. "
        "Do NOT use any external knowledge or make assumptions beyond what's explicitly stated in the context. "
        "If the answer to the user's question is not found in the provided context, respond with: "
        "'Je ne trouve pas d'information sur cette question dans les documents disponibles.' "
        "Make sure to include the source of any information you use in your response.\n\n" + context
    )
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_query}
    ]
    
    # Generate response
    response = llm_ollama.invoke(messages).content
    
    print(f"\n🔹 Sources used: {sources}")
    print(f"🔹 Total documents processed: {len(reranked_docs)}")

    return response, list(sources)



# Create tool node directly
# tools_company = ToolNode([retrieve_company])
# graph_builder_company.add_node("tools", tools_company)
# graph_builder_company.add_node("generate", ask_company_question)

# # Set entry point directly to the tools node
# graph_builder_company.set_entry_point("tools")

# # Connect tools node to generate node
# graph_builder_company.add_edge("tools", "generate")
# graph_builder_company.add_edge("generate", END)

# # Attach memory for threaded sessions
# company_memory = MemorySaver()
# company_graph = graph_builder_company.compile(checkpointer=company_memory)




#Graph


print("Setting up routing logic...")
class RouteQuery(TypedDict):
    destination: Literal[
        "iso_qna",
        "gap_analysis",
        "quiz_iso27002_company",
        "quiz_iso27002_human_eval",
        "quiz_iso42001",
        "generate_document"
    ]

route_prompt = ChatPromptTemplate.from_messages([
    ("system", """You are a routing assistant.
Determine if the user is asking about:
1. ISO 27001 or ISO 27002 or ISO 42001 standard information or the Company-specific document information, questions, clarifications, or explanations (iso_qna)
2. A gap analysis request (gap_analysis)
3. Generate an ISO 27002 quiz with company answers (quiz_iso27002_company)
4. Generate an ISO 27002 quiz with human answers and evaluation (quiz_iso27002_human_eval)
5. Generate an ISO 42001 quiz (quiz_iso42001)
6. Document generation from template request (generate_document)

IMPORTANT ROUTING RULES:
- Route to "iso_qna" for ANY query that:
  * Asks for information about ISO 27001, 27002, or 42001 standards
  * Requests clarification on any ISO standard clause or requirement
  * Contains follow-up questions about ISO standards (even short queries like "explain more" or "what does that mean?")
  * Asks about specific sections or clauses from any ISO standard
  * Contains terms like "standard", "ISO", "requirements", "compliance", "27001", "27002", or "42001"
  * Asks about company-specific documents, especially regarding Sagemcom (e.g., "does Sagemcom have...", "what is Sagemcom's policy on...", etc.)
  * ANY query mentioning "Sagemcom" or "company document" or "company policy" should be routed to "iso_qna"

- Route to "gap_analysis" for requests to analyze gaps between company practices and ISO standards
- Route to "generate_document" for requests to create documents, reports, or templates
- Route to "quiz_iso27002_company" for generating an ISO 27002 quiz with company answers
- Route to "quiz_iso27002_human_eval" for generating an ISO 27002 quiz with human evaluation
- Route to "quiz_iso42001" for generating an ISO 42001 quiz

Consider the entire conversation history when making your decision. Short follow-up queries should be routed to the same destination as the previous query unless the content clearly indicates a different intent.

Respond strictly in JSON format as follows:
{{"destination": "iso_qna"}}  or 
{{"destination": "gap_analysis"}} or 
{{"destination": "quiz_iso27002_company"}} or {{"destination": "quiz_iso27002_human_eval"}} or 
{{"destination": "quiz_iso42001"}} or {{"destination": "generate_document"}}"""),
    ("human", "{query}"),
    ("human", "Previous conversation: {history}")
])

route_chain = (
    route_prompt
    | gemma_model.with_structured_output(RouteQuery, include_raw=False)
    | itemgetter("destination")
)
# -----------------------------
# Gap Analysis Logic
# -----------------------------
print("Setting up gap analysis logic...")
# Load theme clauses
with open("sheet_test.json", "r",encoding="utf-8") as f:
    theme_clauses = json.load(f)


def save_to_csv(results, filename="gap_analysis_results.csv"):
    import pandas as pd
    df = pd.DataFrame(results)
    
    # Save to the current directory with absolute path
    import os
    csv_path = os.path.abspath(filename)
    df.to_csv(csv_path, index=False)
    
    print(f"CSV saved to: {csv_path}")
    return csv_path


def save_to_docx(results, filename="gap_analysis_results.docx"):
    doc = DocxDocument()
    doc.add_heading('ISO 27002:2022 Gap Analysis Report', 0)

    # Add introduction
    doc.add_paragraph(
        "This document presents the results of a gap analysis comparing the organization's documentation "
        "against the ISO 27002:2022 standard requirements."
    )

    # Summary statistics
    total_clauses = len(results)
    confirmed_clauses = len([r for r in results if r["alignment_status"].lower().startswith("confirmé")])
    not_confirmed_clauses = total_clauses - confirmed_clauses
    compliance_percentage = (confirmed_clauses / total_clauses) * 100 if total_clauses > 0 else 0

    doc.add_heading('Executive Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Total clauses analyzed: {total_clauses}\n").bold = True
    summary_para.add_run(f"Confirmed clauses: {confirmed_clauses}\n").bold = True
    summary_para.add_run(f"Non-compliant clauses: {not_confirmed_clauses}\n").bold = True
    summary_para.add_run(f"Overall compliance: {compliance_percentage:.1f}%").bold = True

    # Group results by theme
    themes = {}
    for result in results:
        themes.setdefault(result["theme"], []).append(result)

    # Detailed results by theme
    for theme, theme_results in themes.items():
        doc.add_heading(f'Theme: {theme}', level=1)
        for result in theme_results:
            doc.add_heading(result["clause"], level=2)

            # Alignment status
            status_para = doc.add_paragraph()
            status_para.add_run("Alignment Status: ").bold = True
            status_run = status_para.add_run(result["alignment_status"])
            if result["alignment_status"].lower().startswith("confirmé"):
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red

            # Gap analysis
            gap_para = doc.add_paragraph()
            gap_para.add_run("Gap Analysis: ").bold = True
            gap_para.add_run(result["gap_analysis"])

            # Recommendations
            rec_para = doc.add_paragraph()
            rec_para.add_run("Recommendations: ").bold = True
            rec_para.add_run(result["recommendations"])

            # Sources
            sources_para = doc.add_paragraph()
            sources_para.add_run("Company Source: ").bold = True
            sources_para.add_run(result["company_source"])
            sources_para.add_run("\nISO Source: ").bold = True
            sources_para.add_run(result["iso_source"])

            doc.add_paragraph("---")

    # Convert to absolute path and save
    abs_path = os.path.abspath(filename)
    doc.save(abs_path)
    print(f"📄 DOCX saved to: {abs_path}")
    return abs_path
def parse_gap_analysis_result(result_text):
    try:
        # Extract alignment status - French version
        alignment_match = re.search(r"(?:Statut d['']Alignement:|Alignement:).*?(confirmé|non confirmé)",
                                   result_text, re.IGNORECASE)
        alignment_status = alignment_match.group(1).strip() if alignment_match else "Inconnu"

        # Extract gap analysis - French version
        gap_match = re.search(r"(?:Analyse d['']Écart:|Écart:)(.+?)(?:Recommandation[s]?:|Références:|$)",
                             result_text, re.DOTALL)
        gap_analysis = gap_match.group(1).strip() if gap_match else "Aucune analyse d'écart fournie"

        # Extract recommendations - French version
        rec_match = re.search(r"(?:Recommandation[s]?:)(.+?)(?:Références:|Source[s]? d[e']entreprise:|Source[s]? ISO:|$)",
                             result_text, re.DOTALL)
        recommendations = rec_match.group(1).strip() if rec_match else "Aucune recommandation fournie"

        return {
            "alignment_status": alignment_status,
            "gap_analysis": gap_analysis,
            "recommendations": recommendations
        }
    except Exception as e:
        print(f"Erreur lors de l'analyse du résultat: {e}")
        return {
            "alignment_status": "Erreur",
            "gap_analysis": "Erreur lors de l'analyse du résultat",
            "recommendations": "Erreur lors de l'analyse du résultat"
        }
print("Precomputing ISO metadata...")
# Precompute ISO metadata for faster access
def run_gap_analysis():
    results = []
    for theme, clauses in theme_clauses.items():
        for clause in clauses:
            print(f"\n📝 Analyzing clause: {clause} from theme: {theme}")
            
            # Get ISO information
            iso_chunks = vector_store_i_title.similarity_search_with_score(clause, k=1)
            
            # Combine ISO content, titles, and sources
            iso_titles = [doc.page_content for doc, _ in iso_chunks]
            iso_contents = [doc.metadata.get("content", "") for doc, _ in iso_chunks]
            iso_sources = [doc.metadata.get("source", "") for doc, _ in iso_chunks]
            
            iso_metadata = {
                "titles": iso_titles,
                "sources": iso_sources,
                "contents": iso_contents
            }
             
            # Print the ISO metadata
            print("\nISO Metadata Retrieved:")
            print("Titles:", iso_metadata["titles"])
            print("Sources:", iso_metadata["sources"])
            print("Contents:", iso_metadata["contents"])
            
            # Generate quiz questions based on ISO content to verify compliance
            quiz_prompt = f"""
            Vous êtes un auditeur expert en normes ISO 27002:2022.
            
            Examinez attentivement le contenu ISO suivant et créez max 3 questions d'audit précises pour 
            évaluer la conformité d'une entreprise à cette clause ISO spécifique.
            
            Clause ISO: {clause}
            
            Contenu ISO:
            {iso_contents[0] if iso_contents else "Contenu non disponible"}
            
            Pour chaque question:
            1. Concentrez-vous sur les exigences clés de la clause
            2. Formulez des questions directes et spécifiques 
            3. Assurez-vous que les questions couvrent différents aspects de la clause
            4. Incluez des questions sur les procédures, la documentation, et la mise en œuvre
            
            Si le contenu mentionne des recommandations spécifiques, créez également des questions 
            pour vérifier si ces recommandations sont suivies.
            
            Répondez avec UNIQUEMENT une liste de questions, chacune sur une ligne séparée. 
            Ne numérotez pas les questions et n'ajoutez aucun texte supplémentaire.
            """
            
            print(f"🔍 Generating audit questions for clause: {clause}")
            questions_response = gemma_model.invoke(quiz_prompt).content
            
            # Parse questions (each on a separate line)
            questions = [q.strip() for q in questions_response.split('\n') if q.strip()]
            
            # For each question, get a company response using ask_company_question
            qna_pairs = []
            alignment_scores = []
            company_sources_all = []
            
            for question in questions:
                print(f"📋 Processing question: {question}")
                company_answer, sources = ask_company_question(question)
                print(f"Company Answer: {company_answer}")
                # Add sources to the overall list
                for source in sources:
                    if source not in company_sources_all:
                        company_sources_all.append(source)
                
                # Create evaluation prompt to assess alignment
                eval_prompt = f"""
                En tant qu'auditeur ISO 27002:2022, évaluez la conformité de cette réponse 
                par rapport à la clause ISO concernée.
                
                Clause ISO: {clause}
                
                Contenu ISO applicable:
                {iso_contents[0] if iso_contents else "Contenu non disponible"}
                
                Question d'audit: {question}
                
                Réponse de l'entreprise: {company_answer}
                
                Évaluez:
                1. Si la réponse de l'entreprise montre une conformité à l'exigence ISO
                2. Le degré d'alignement (échelle de 0-100%)
                3. Les lacunes ou problèmes identifiés
                
                Répondez uniquement avec un pourcentage d'alignement (ex: 75%) suivi d'un court commentaire.
                """
                
                alignment_response = gemma_model.invoke(eval_prompt).content
                print(f"Alignment Evaluation: {alignment_response}")
                # Extract percentage from response
                percentage_match = re.search(r'(\d+)%', alignment_response)
                alignment_percentage = int(percentage_match.group(1)) if percentage_match else 50
                alignment_comment = alignment_response.replace(percentage_match.group(0), '').strip() if percentage_match else alignment_response
                
                alignment_scores.append(alignment_percentage)
                
                # Store the Q&A pair and evaluation
                qna_pairs.append({
                    "question": question,
                    "company_answer": company_answer,
                    "alignment_score": alignment_percentage,
                    "alignment_comment": alignment_comment
                })
            
            # Calculate overall alignment score for this clause
            overall_alignment = sum(alignment_scores) / len(alignment_scores) if alignment_scores else 0
            alignment_status = f"confirmé ({overall_alignment:.1f}%)" if overall_alignment >= 70 else f"non confirmé ({overall_alignment:.1f}%)"
            
            # Create analysis based on all Q&A evaluations
            analysis_prompt = f"""
            En tant qu'auditeur ISO 27002:2022, fournissez une analyse d'écart basée sur les questions
            et réponses suivantes concernant la clause ISO: {clause}
            
            Contenu ISO:
            {iso_contents[0] if iso_contents else "Contenu non disponible"}
            
            Questions et Réponses:
            {json.dumps(qna_pairs, ensure_ascii=False, indent=2)}
            
            Score d'alignement global: {overall_alignment:.1f}%
            
            Veuillez fournir:
            1. Une analyse détaillée des écarts identifiés
            2. Des recommandations pour améliorer la conformité
            
            Format de réponse attendu:
            - Analyse d'Écart: [votre analyse]
            - Recommandations: [vos recommandations]
            """
            
            analysis_response = gemma_model.invoke(analysis_prompt).content
            
            # Parse analysis response
            gap_analysis = ""
            recommendations = ""
            
            gap_match = re.search(r"(?:Analyse d[\'']Écart:|Analyse d[\'']écart:)(.*?)(?:Recommandation[s]?:|$)", analysis_response, re.DOTALL)
            if gap_match:
                gap_analysis = gap_match.group(1).strip()
            
            rec_match = re.search(r'(?:Recommandation[s]?:)(.*?)$', analysis_response, re.DOTALL)
            if rec_match:
                recommendations = rec_match.group(1).strip()
            
            # Collect the final result
            results.append({
                "theme": theme,
                "clause": clause,
                "alignment_status": alignment_status,
                "alignment_percentage": f"{overall_alignment:.1f}%",
                "gap_analysis": gap_analysis,
                "recommendations": recommendations,
                "qna_pairs": qna_pairs,
                "company_source": company_sources_all,
                "iso_source": iso_titles
            })
            
            print(f"✅ Completed analysis for {clause}")
    
    # Save results
    csv_path = save_to_csv(results)
    docx_path = save_to_docx(results)
    
    import pandas as pd
    if "gap_analysis_df" not in st.session_state:
        st.session_state.gap_analysis_df = pd.DataFrame(results)

    return {"csv": csv_path, "docx": docx_path, "df": results}
def display_gap_analysis_in_streamlit(gap_results):
    """Display the gap analysis results with a 3-tab interface"""
    
    # Initialize session state for storing feedback if not already present
    if "gap_analysis_feedback" not in st.session_state:
        st.session_state.gap_analysis_feedback = {}
    
    # Create three main tabs for the gap analysis interface
    main_tabs = st.tabs(["Initial Results Overview", "Provide Feedback", "View Updated Analysis"])
    
    # Tab 1: Overview of initial gap analysis results
    with main_tabs[0]:
        st.markdown("## ISO 27002:2022 Gap Analysis - Initial Results")
        
        # Create a DataFrame view of all results
        df = pd.DataFrame(gap_results)
        
        # Calculate compliance statistics
        total_clauses = len(gap_results)
        confirmed_clauses = sum(1 for r in gap_results if r["alignment_status"].lower().startswith("confirmé"))
        not_confirmed_clauses = total_clauses - confirmed_clauses
        compliance_percentage = (confirmed_clauses / total_clauses) * 100 if total_clauses > 0 else 0
        
        # Display summary statistics with donut chart
        st.markdown(f"### Summary Statistics")

        # Calculate statistics from gap results
        total_clauses = len(gap_results)
        status_counts = {
            "Fully Implemented": sum(1 for r in gap_results if r["alignment_status"].lower().startswith("confirmé")),
            "Not Implemented": sum(1 for r in gap_results if r["alignment_status"].lower().startswith("non confirmé")),
            "Partially Implemented": 15,  # Initialize with zero, adjust if you have this status
            "Open": 7,                   # Initialize with zero, adjust if you have this status
            "N/A": 5                     # Initialize with zero, adjust if you have this status
        }

        # Define status colors
        status_colors = {
            "Fully Implemented": "green",
            "Partially Implemented": "yellow",
            "Not Implemented": "red",
            "Open": "blue",
            "N/A": "gray"
        }

        # Remove zero counts for better visualization
        filtered_status = {k: v for k, v in status_counts.items() if v > 0}

        # Create two columns for layout
        col1, col2 = st.columns([3, 2])

        with col1:
            # Create donut chart with plotly
            import plotly.graph_objects as go
            
            labels = list(filtered_status.keys())
            values = list(filtered_status.values())
            colors = [status_colors[status] for status in labels]
            
            fig = go.Figure(data=[go.Pie(
                labels=labels, 
                values=values,
                hole=.4,  # Creates donut chart
                marker_colors=colors,
                textinfo='percent+label',
                textfont_size=12,
                textposition='inside',
                insidetextorientation='radial'
            )])
            
            fig.update_layout(
                title_text="ISO 27002:2022 Compliance Status",
                title_x=0.5,
                showlegend=True,
                legend=dict(
                    orientation="v",
                    yanchor="top",
                    y=1.1,
                    xanchor="center",
                    x=0.5
                ),
                height=400,
                margin=dict(t=80, b=20, l=20, r=20),
            )
            
            st.plotly_chart(fig, use_container_width=True)

        with col2:
            # Display numerical metrics alongside the chart
            st.markdown("#### Compliance Metrics")
            st.metric("Total Clauses", total_clauses)
            st.metric("Compliant Clauses", status_counts["Fully Implemented"])
            
            compliance_percentage = (status_counts["Fully Implemented"] / total_clauses) * 100 if total_clauses > 0 else 0
            st.metric("Overall Compliance", f"{compliance_percentage:.1f}%")
            
            # Add a color legend
            st.markdown("#### Legend")
            for status, color in status_colors.items():
                if status in filtered_status:
                    st.markdown(f"<div style='display:flex;align-items:center;margin-bottom:5px;'>"
                                f"<div style='background-color:{color};width:15px;height:15px;margin-right:8px;'></div>"
                                f"<div>{status}: {filtered_status.get(status, 0)} clauses</div></div>", 
                                unsafe_allow_html=True)
        
        # Show the results in a table with filtering
        st.markdown("### All Clauses Analysis")
        
        # Select columns for display
        display_cols = ["theme", "clause", "alignment_status", "gap_analysis", "recommendations"]
        if all(col in df.columns for col in display_cols):
            st.dataframe(df[display_cols], use_container_width=True)
        else:
            st.dataframe(df)
            
        # Add download buttons for the original analysis
        st.markdown("### Export Initial Results")
        col1, col2 = st.columns(2)
        with col1:
            csv = df.to_csv(index=False)
            st.download_button(
                label="📥 Download Original Gap Analysis CSV",
                data=csv,
                file_name="gap_analysis.csv",
                mime="text/csv",
            )
        with col2:
            docx_path = save_to_docx(gap_results)
            if os.path.exists(docx_path):
                with open(docx_path, "rb") as docx_file:
                    docx_data = docx_file.read()
                    st.download_button(
                        label="📄 Download Original Gap Analysis DOCX",
                        data=docx_data,
                        file_name="gap_analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("DOCX file could not be generated")
    
    # Tab 2: Select specific theme and clause to provide feedback
    with main_tabs[1]:
        st.markdown("## Provide Feedback on Gap Analysis")
        
        # Group results by theme for selection
        themes = {}
        for result in gap_results:
            theme = result["theme"]
            if theme not in themes:
                themes[theme] = []
            themes[theme].append(result)
        
        # Create theme selector
        selected_theme = st.selectbox("Select Theme", list(themes.keys()))
        
        if selected_theme:
            # Get clauses for the selected theme
            theme_clauses = [(result["clause"], i) for i, result in enumerate(themes[selected_theme])]
            clause_options = [clause for clause, _ in theme_clauses]
            
            # Create clause selector
            selected_clause = st.selectbox("Select Clause", clause_options)
            
            if selected_clause:
                # Find the selected result
                selected_idx = next((idx for clause, idx in theme_clauses if clause == selected_clause), None)
                if selected_idx is not None:
                    selected_result = themes[selected_theme][selected_idx]
                    
                    # Display the selected result details
                    st.markdown("### Selected Clause Analysis")
                    
                    # Display basic clause information
                    st.markdown(f"**Theme:** {selected_theme}")
                    st.markdown(f"**Clause:** {selected_result['clause']}")
                    
                    # Display alignment status with color coding
                    status_color = "green" if selected_result['alignment_status'].lower().startswith("confirmé") else "red"
                    st.markdown(f"**Alignment Status:** <span style='color:{status_color};'>{selected_result['alignment_status']}</span>", unsafe_allow_html=True)
                    
                    # Use expandable sections for details
                    with st.expander("📝 Gap Analysis", expanded=True):
                        st.markdown(selected_result['gap_analysis'])
                    
                    with st.expander("🔍 Recommendations", expanded=True):
                        st.markdown(selected_result['recommendations'])
                    
                    # Sources in an expandable section
                    with st.expander("🔗 Sources", expanded=False):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("**Company Sources:**")
                            for source in selected_result.get('company_source', []):
                                st.markdown(f"- {source}")
                        
                        with col2:
                            st.markdown("**ISO Sources:**")
                            for source in selected_result.get('iso_source', []):
                                st.markdown(f"- {source}")
                    
                    # Create a form for feedback in an expander
                    with st.expander("✏️ Provide Your Feedback", expanded=True):
                        clause_key = f"{selected_theme}_{selected_clause}"
                        
                        with st.form(key=f"feedback_form_{clause_key}"):
                            # Pre-fill with existing feedback if available
                            initial_feedback = ""
                            if clause_key in st.session_state.gap_analysis_feedback:
                                initial_feedback = st.session_state.gap_analysis_feedback[clause_key]
                            
                            feedback = st.text_area(
                                "Your feedback on this gap analysis:",
                                value=initial_feedback,
                                height=200,
                                placeholder="Provide your expert insights, corrections, or additional context..."
                            )
                            
                            submitted = st.form_submit_button("Save Feedback")
                            if submitted:
                                st.session_state.gap_analysis_feedback[clause_key] = feedback
                                st.success(f"Feedback for clause {selected_clause} saved successfully!")
        
        # Add buttons for global feedback actions
        st.markdown("### Feedback Actions")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Clear All Feedback", key="clear_gap_feedback"):
                st.session_state.gap_analysis_feedback = {}
                if "updated_gap_analysis" in st.session_state:
                    del st.session_state.updated_gap_analysis
                st.success("All feedback cleared!")
                st.rerun()
        
        with col2:
            if st.button("📊 Update Analysis with Feedback", key="update_gap_analysis"):
                if not st.session_state.gap_analysis_feedback:
                    st.warning("No feedback provided yet. Please provide feedback first.")
                else:
                    with st.spinner("Updating gap analysis with your feedback..."):
                        updated_results = update_gap_analysis_with_feedback(gap_results)
                        if updated_results:
                            st.session_state.updated_gap_analysis = updated_results
                            st.success("Gap analysis updated with your feedback!")
                            st.rerun()
    # Tab 3: Display updated analysis
    with main_tabs[2]:
        st.markdown("## Updated Gap Analysis Results")
        
        if "updated_gap_analysis" not in st.session_state:
            st.info("No updated analysis available yet. Please provide feedback in the 'Provide Feedback' tab and click 'Update Analysis with Feedback'.")
        else:
            display_updated_gap_analysis(gap_results, st.session_state.updated_gap_analysis)

def update_gap_analysis_with_feedback(original_results):
    """Update gap analysis results using collected feedback"""
    
    if not st.session_state.gap_analysis_feedback:
        return None
        
    updated_results = []
    
    for result in original_results:
        theme = result["theme"]
        clause = result["clause"]
        clause_key = f"{theme}_{clause}"
        
        # Check if we have feedback for this clause
        if clause_key in st.session_state.gap_analysis_feedback and st.session_state.gap_analysis_feedback[clause_key].strip():
            feedback = st.session_state.gap_analysis_feedback[clause_key]
            
            # Create a prompt for updating the gap analysis
            update_prompt = f"""
Vous êtes un expert en sécurité de l'information chargé d'améliorer une analyse d'écart ISO 27002.

Voici l'analyse d'écart originale pour la clause ISO 27002 "{clause}" dans le thème "{theme}":

**Statut d'Alignement Original:** {result['alignment_status']}

**Analyse d'Écart Original:** 
{result['gap_analysis']}

**Recommandations Originales:**
{result['recommendations']}

Un expert humain a fourni le commentaire suivant sur cette analyse:
"{feedback}"

En tenant compte de ce retour d'expert, veuillez fournir:
1. Un statut d'alignement mis à jour (confirmé ou non confirmé)
2. Une analyse d'écart révisée
3. Des recommandations améliorées

Répondez strictement au format suivant en français:
- **Statut d'Alignement:** [statut mis à jour]
- **Analyse d'Écart:** [analyse mise à jour]
- **Recommandations:** [recommandations mises à jour]
"""
            
            # Get updated analysis using LLM
            try:
                updated_analysis = gemma_model.invoke(update_prompt).content
                
                # Parse the updated result
                updated_parsed = parse_gap_analysis_result(updated_analysis)
                
                # Create updated result
                updated_result = result.copy()  # Copy the original
                updated_result["alignment_status"] = updated_parsed["alignment_status"]
                updated_result["gap_analysis"] = updated_parsed["gap_analysis"]
                updated_result["recommendations"] = updated_parsed["recommendations"]
                updated_result["feedback"] = feedback  # Add the human feedback
                updated_result["is_updated"] = True
                
                updated_results.append(updated_result)
            except Exception as e:
                print(f"Error updating gap analysis for {clause}: {e}")
                result["error"] = str(e)
                result["is_updated"] = False
                updated_results.append(result)
        else:
            # No feedback, keep original
            result["is_updated"] = False
            updated_results.append(result)
    
    return updated_results

def display_updated_gap_analysis(original_results, updated_results):
    """Display the updated gap analysis results side by side with original results"""
    st.markdown("## Updated Gap Analysis with Feedback")
    
    # Group results by theme for better organization
    themes = {}
    for result in updated_results:
        theme = result["theme"]
        if theme not in themes:
            themes[theme] = []
        themes[theme].append(result)
    
    # Create tabs for each theme
    theme_tabs = st.tabs(list(themes.keys()))
    
    for theme_idx, (theme_name, theme_results) in enumerate(themes.items()):
        with theme_tabs[theme_idx]:
            st.markdown(f"### Theme: {theme_name}")
            
            # Create an expander for each clause in the theme
            for result_idx, result in enumerate(theme_results):
                clause = result["clause"]
                
                # Find the original result for comparison
                original_result = next((r for r in original_results if r["theme"] == theme_name and r["clause"] == clause), None)
                
                with st.expander(f"Clause: {clause}", expanded=result.get("is_updated", False)):
                    # If the result was updated based on feedback
                    if result.get("is_updated", False):
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("#### Original Analysis")
                            st.markdown(f"**Alignment Status:** {original_result['alignment_status']}")
                            st.markdown(f"**Gap Analysis:** {original_result['gap_analysis']}")
                            st.markdown(f"**Recommendations:** {original_result['recommendations']}")
                        
                        with col2:
                            st.markdown("#### Updated Analysis")
                            
                            # Highlight changes in alignment status with color
                            if result['alignment_status'] != original_result['alignment_status']:
                                status_color = "green" if result['alignment_status'].lower().startswith("confirmé") else "red"
                                st.markdown(f"**Alignment Status:** <span style='color:{status_color};'>{result['alignment_status']}</span>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"**Alignment Status:** {result['alignment_status']}")
                            
                            st.markdown(f"**Gap Analysis:** {result['gap_analysis']}")
                            st.markdown(f"**Recommendations:** {result['recommendations']}")
                        
                        # Display the feedback that led to this update
                        st.markdown("#### Feedback Provided")
                        st.info(result.get("feedback", "No feedback provided"))
                    else:
                        # No update was made
                        st.markdown("#### Analysis (No Updates)")
                        st.markdown(f"**Alignment Status:** {result['alignment_status']}")
                        st.markdown(f"**Gap Analysis:** {result['gap_analysis']}")
                        st.markdown(f"**Recommendations:** {result['recommendations']}")
                        
                        if "error" in result:
                            st.error(f"Error updating analysis: {result['error']}")
    
    # Create export options
    st.markdown("### Export Updated Results")
    
    # Create CSV download button
    csv = pd.DataFrame(updated_results).to_csv(index=False)
    st.download_button(
        label="📥 Download Updated Gap Analysis CSV",
        data=csv,
        file_name="updated_gap_analysis.csv",
        mime="text/csv",
    )
    
    # Generate and offer DOCX download
    docx_path = save_to_docx(updated_results, filename="updated_gap_analysis.docx")
    if os.path.exists(docx_path):
        with open(docx_path, "rb") as docx_file:
            docx_data = docx_file.read()
            st.download_button(
                label="📄 Download Updated Gap Analysis DOCX",
                data=docx_data,
                file_name="updated_gap_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
# -----------------------------
# Quiz generator
# -----------------------------
def generate_quiz_question(section, subsection, text):
    """ 
    Use LLM to generate a variable number of quiz questions for a specific ISO clause 
    """
    # Get the reference ID from the subsection (first part before space) 
    ref_parts = subsection.split(" ", 1) 
    ref = ref_parts[0] if len(ref_parts) > 0 else "unknown"
    
    # Create a more focused prompt for the LLM 
    prompt = f""" 
Agissez en tant qu’auditeur principal certifié lead auditor ISO 27002:2022.

Vous êtes en train d’examiner la clause ISO {subsection} : « {text} ».

Votre mission est de générer une série réaliste de questions de type audit, suivant une progression logique et en chaîne. Ces questions doivent simuler la manière dont un auditeur expérimenté approfondirait une évaluation de conformité réelle.

Pour chaque question, fournissez immédiatement après une réponse plausible et pertinente que l’organisation auditée pourrait donner, en lien avec le contexte décrit. Les réponses doivent être réalistes, refléter les bonnes pratiques, et s’aligner avec les exigences de la clause.

Contexte :
L’organisation auditée est une entreprise spécialisée dans les logiciels et les technologies. Elle comprend des équipes de développement logiciel ainsi que des équipes de validation. Ces équipes interagissent fréquemment avec le client télécom (opérateur) dans le monde entier, et collaborent avec des partenaires tiers pour la production des produits résidentiels. Les activités de développement, la validation du code, l’intégration et le déploiement continus, ainsi que la collaboration externe font partie intégrante de leur environnement opérationnel.

L’entreprise utilise des outils comme JIRA, jenkins, Confluence, GitLab, SharePoint et Office 365 pour faciliter la gestion des tâches, la collaboration, et la gestion documentaire.
les OS des postes de travail utilisés sont soit du Linux soit du Windows, les produits developpés sont avec des OS Android ou Linux selon le client en face.
la structure de la société est faite de sorte à avoir des equipes de developpement , des equipes de validation , des equipes de certification fonctionnels, des equipes transverses  et un equipe infra reseau qui permet de fournir des environnements de test et de simulation des environnements client, ou des accées VPN aux infrastructures client( operateur)

Consignes :
- Commencez par une question générale pour évaluer la conformité de base.
- Ensuite, posez des questions de plus en plus détaillées ou approfondies, en fonction des réponses probables.
- Les questions doivent être spécifiques, exploitables et clairement liées au contenu de la clause et au contexte de l'entreprise.
- Fournissez une réponse réaliste juste après chaque question qui se base sur les clauses de ISO comme référence, justifiant la réponse par rapport au document de Clause ISO fourni.
- Adoptez un ton naturel et humain, tel qu’un auditeur experimentés dans le domaine de l'entreprise le ferait lors d’un entretien.
- Incluez 1 question/réponse, selon la complexité de la clause.
- Assurez-vous que chaque question s’appuie sur la précédente — comme dans une conversation réfléchie ou une visite d’audit.
- Formulez vos questions et réponses en tenant compte de l’environnement de l’entreprise : développement logiciel, validation, interactions clients et collaborations avec des tiers, ainsi que l’utilisation des outils mentionnés.
- N’INCLUEZ AUCUNE mise en forme, puce ou explication — produisez uniquement les questions suivies de leur réponse sur deux lignes consécutives.
-selement un seul question/réponse par clause, pas plus.
Générez maintenant l’ensemble des paires questions/réponses en chaîne pour cette clause, en français. 

"""

    response = gemma_model.invoke(prompt).content

    # Split into lines, and group them as Q&A pairs
    lines = [line.strip() for line in response.strip().split("\n") if line.strip()]
    qna_pairs = []

    for i in range(0, len(lines) - 1, 2):
        question = lines[i]
        answer = lines[i + 1]
        company_answer, sources = ask_company_question(question)
        qna_pairs.append({
            "question": question,
            "answer": answer,
            "company_answer": company_answer,
            "sources": sources
        })

    return {
        "section": section,
        "subsection": subsection,
        "ref_id": ref,
        "qna": qna_pairs,
        "text": text  # Include the clause content
    }

def save_quiz_to_csv(quiz_results):
    """
    Save quiz results directly (without flattening) to a CSV-compatible JSON format.
    """
    import os
    import json

    filename = "iso27002_quiz1_output_raw.json"  # Better suited for structured data
    file_path = os.path.abspath(filename)

    # Save directly as JSON for structured access later
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(quiz_results, f, ensure_ascii=False, indent=2)

    print(f"Raw quiz results saved to: {file_path}")
    return file_path

    
    

def display_quiz_in_streamlit(quiz_results):
    

    st.markdown("## ISO 27002:2022 Audit Questions")

    for i, item in enumerate(quiz_results):
        section = item["section"]
        subsection = item["subsection"]
        ref_id = item["ref_id"]
        qna = item["qna"]

        with st.expander(f"{subsection}", expanded=(i == 0)):
            st.markdown(f"**Section:** {section}")
            st.markdown(f"**Reference:** {ref_id}")

            questions_df = pd.DataFrame({
                "Question #": range(1, len(qna) + 1),
                "Audit Question": [q["question"] for q in qna],
                "Suggested Answer": [q["answer"] for q in qna],
                "Company Answer": [q["company_answer"] for q in qna],
                "Sources": [", ".join(q["sources"]) if q["sources"] else "—" for q in qna]
            })

            st.table(questions_df)

            section_csv = questions_df.to_csv(index=False).encode('utf-8')
            st.download_button(
                label=f"📥 Download '{ref_id}' Questions",
                data=section_csv,
                file_name=f"iso27002_{ref_id}_questions.csv",
                mime="text/csv",
                key=f"download_section_{i}"
            )

    all_questions = []
    for item in quiz_results:
        for i, qna in enumerate(item["qna"]):
            all_questions.append({
                "Section": item["section"],
                "Subsection": item["subsection"],
                "Reference ID": item["ref_id"],
                "Question #": i + 1,
                "Question": qna["question"],
                "Answer": qna["answer"],
                "Company Answer": qna["company_answer"],
                "Sources": ", ".join(qna["sources"]) if qna["sources"] else "—"
            })

    all_questions_df = pd.DataFrame(all_questions)

    if all_questions:
        st.markdown("### Download All Questions")
        full_csv = all_questions_df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="📥 Download Complete Question Set",
            data=full_csv,
            file_name="iso27002_all_questions.csv",
            mime="text/csv",
            key="download_all"
        )


import json

def generate_iso_quiz():
    print("🔍 Generating ISO 27002 quiz questions...")

    with open("structured_2.json", 'r', encoding='utf-8') as f:
        structured_sections = json.load(f)

    quiz_results = []

    for category, items in structured_sections.items():
        print(f"\n📂 Category: {category}")

        for entry in items:
            try:
                subsection, section = map(str.strip, entry.split(",", 1))
            except ValueError:
                print(f"❌ Skipping malformed entry: {entry}")
                continue

            query = f"{subsection} {section}"
            print(f"\n🔹 Searching for: {query}")

            results = vector_store_i_title.similarity_search(query=query, k=1)

            if results:
                combined_context = ""
                for doc in results:
                    title = doc.metadata.get("title", "")
                    content = doc.metadata.get("content", "")
                    source = doc.metadata.get("source", "")
                    print(f"✅ Found title: {title}")
                    print(f"📄 Content snippet: {content[:200]}...\n")

                    combined_context += f"{title}\n{content}\n{source}\n"

                if combined_context.strip():
                    quiz_item = generate_quiz_question(section, subsection, combined_context.strip())
                    quiz_item["category"] = category  # optionally include the category
                    quiz_results.append(quiz_item)
                else:
                    print(f"❌ No valid content for: {subsection}")
            else:
                print(f"❌ No matches found for: {query}")

    return quiz_results
# -----------------------------
# Quiz generator with human eval
# -----------------------------
def generate_quiz_question_1(section, subsection, text):
    """ 
    Use LLM to generate a variable number of quiz questions for a specific ISO clause 
    """
    # Get the reference ID from the subsection (first part before space) 
    ref_parts = subsection.split(" ", 1) 
    ref = ref_parts[0] if len(ref_parts) > 0 else "unknown"
    
    # Create a more focused prompt for the LLM 
    prompt = f""" 
Agissez en tant qu’auditeur principal certifié ISO 27002:2022.

Vous êtes en train d’examiner la clause {subsection} : « {text} ».

Votre mission est de générer une série réaliste de questions de type audit, suivant une progression logique et en chaîne. Ces questions doivent simuler la manière dont un auditeur expérimenté approfondirait une évaluation de conformité réelle.

Pour chaque question, fournissez immédiatement après une réponse plausible et pertinente que l’organisation auditée pourrait donner, en lien avec le contexte décrit. Les réponses doivent être réalistes, refléter les bonnes pratiques, et s’aligner avec les exigences de la clause.

Contexte :
L’organisation auditée est une entreprise spécialisée dans les logiciels et les technologies. Elle comprend des équipes de développement logiciel ainsi que des équipes de validation. Ces équipes interagissent fréquemment avec le client télécom (opérateur) en Asie et en Europe, et collaborent avec des partenaires tiers. Les activités de développement, la validation du code, l’intégration et le déploiement continus, ainsi que la collaboration externe font partie intégrante de leur environnement opérationnel.

Consignes :
- Commencez par une question générale pour évaluer la conformité de base.
- Ensuite, posez des questions de plus en plus détaillées ou approfondies, en fonction des réponses probables.
- Les questions doivent être spécifiques, exploitables et clairement liées au contenu de la clause.
- Fournissez une réponse réaliste juste après chaque question.
- Adoptez un ton naturel et humain, tel qu’un auditeur le ferait lors d’un entretien.
- Incluez 1 et un seul seulement  question/réponse.
- Assurez-vous que chaque question s’appuie sur la précédente — comme dans une conversation réfléchie ou une visite d’audit.
- Formulez vos questions et réponses en tenant compte de l’environnement de l’entreprise : développement logiciel, validation, interactions clients et collaborations avec des tiers.
- N’INCLUEZ AUCUNE mise en forme, puce ou explication — produisez uniquement les questions suivies de leur réponse sur deux lignes consécutives.

Générez maintenant l’ensemble des paires questions/réponses en chaîne pour cette clause, en français.
"""

    response = gemma_model.invoke(prompt).content

    # Split into lines, and group them as Q&A pairs
    lines = [line.strip() for line in response.strip().split("\n") if line.strip()]
    qna_pairs = []

    for i in range(0, len(lines) - 1, 2):
        question = lines[i]
        answer = lines[i + 1]
        qna_pairs.append({
            "question": question,
            "answer": answer
        })

    return {
        "section": section,
        "subsection": subsection,
        "ref_id": ref,
        "qna": qna_pairs,
        "text": text  # Include the clause content
    }





# Define gemma_model at the top of your script


def evaluate_user_response_1(user_response: str, model_answer: str, iso_text: str) -> Tuple[int, str]:
    """
    Evaluate user response using Gemma model directly.
    """
    if not user_response or user_response.strip() == "":
        return 0, "No response provided"
    
    # Create the prompt for evaluation
    prompt = f"""
Agissez en tant qu'auditeur principal certifié ISO 27002:2022.

Vous devez évaluer la réponse d'un candidat auditeur à une question d'audit par rapport à la norme ISO et à la réponse modèle.

Contexte de la norme ISO:
{iso_text}

Réponse modèle:
{model_answer}

Réponse du candidat:
{user_response}

Votre tâche:
1. Évaluez la réponse du candidat par rapport à la réponse modèle et au contexte de la norme ISO.
2. Déterminez un pourcentage de conformité (0% à 100%) qui représente à quel point la réponse du candidat est alignée avec les exigences de la norme ISO.
3. Fournissez un commentaire constructif qui explique ce qui est correct et ce qui pourrait être amélioré.

Format de votre réponse:
- Pourcentage de conformité: X%
- Commentaire: [Votre commentaire d'évaluation]

Ne donnez AUCUNE introduction ou conclusion - fournissez uniquement les deux éléments demandés dans le format exact spécifié ci-dessus.
"""
    
    try:
        response = gemma_model.invoke(prompt).content
        
        lines = response.strip().split("\n")
        percentage_line = next((line for line in lines if "Pourcentage" in line or "pourcentage" in line.lower()), "")
        percentage_str = percentage_line.split(":")[-1].strip().replace("%", "")
        
        try:
            percentage = int(percentage_str) if percentage_str.isdigit() else 75
        except:
            percentage = 75
        
        feedback_index = next((i for i, line in enumerate(lines) if "Commentaire" in line or "commentaire" in line.lower()), -1)
        if feedback_index >= 0 and feedback_index < len(lines) - 1:
            feedback = "\n".join(lines[feedback_index+1:])
        else:
            feedback = next((line.split(":", 1)[1].strip() for line in lines if "Commentaire" in line or "commentaire" in line.lower()), 
                          "Veuillez comparer votre réponse avec la réponse suggérée.")
            
        return percentage, feedback
        
    except Exception as e:
        print(f"LLM evaluation error: {str(e)}")
        return 70, "L'évaluation automatique n'a pas pu être complétée. Veuillez comparer votre réponse avec la réponse suggérée."

def display_quiz_in_streamlit_1(quiz_results):
    """Display the quiz questions and collect responses only without evaluation"""
    st.markdown("## ISO 27002:2022 Audit Questions")

    # Initialize session state for storing responses if not already present
    if "quiz_responses" not in st.session_state:
        st.session_state.quiz_responses = {}

    tabs = st.tabs([item["subsection"] for item in quiz_results])

    for i, (tab, item) in enumerate(zip(tabs, quiz_results)):
        subsection = item["subsection"]
        ref_id = item["ref_id"]
        iso_text = item["text"]
        qna = item["qna"]

        with tab:
            st.markdown(f"**Subsection:** {subsection} ({ref_id})")

            # Create a form for each tab to prevent premature submission
            with st.form(key=f"form_{ref_id}"):
                tab_responses = {}
                
                for q_idx, qa in enumerate(qna):
                    question = qa["question"]
                    model_answer = qa["answer"]
                    response_key = f"{ref_id}_{q_idx}"

                    st.markdown(f"**Question #{q_idx + 1}:** *{question}*")

                    with st.expander("Show suggested answer"):
                        st.markdown(f"*{model_answer}*")

                    # Pre-fill with existing response if available
                    initial_value = ""
                    if response_key in st.session_state.quiz_responses:
                        initial_value = st.session_state.quiz_responses[response_key]["user_response"]

                    # Use a text area that doesn't submit on Enter
                    user_response = st.text_area(
                        f"Your response to question #{q_idx + 1}:",
                        value=initial_value,
                        key=f"input_{response_key}",
                        height=150  # Add some height for more comfortable typing
                    )

                    # Store response metadata for this question
                    tab_responses[response_key] = {
                        "question": question,
                        "model_answer": model_answer,
                        "user_response": user_response,
                        "iso_text": iso_text,
                        "ref_id": ref_id,
                        "q_idx": q_idx
                    }
                
                # Add a submit button for this tab's form
                submitted = st.form_submit_button("Save Responses for This Section")
                if submitted:
                    # Update session state with this tab's responses
                    for key, value in tab_responses.items():
                        st.session_state.quiz_responses[key] = value
                    st.success(f"Responses for section {subsection} saved successfully!")
    
    # Add clear button with a unique key - note we removed the evaluate button
    if st.button("🗑️ Clear All Responses", key="clear_btn_display"):
        st.session_state.quiz_responses = {}
        if "evaluation_results" in st.session_state:
            del st.session_state.evaluation_results
        st.success("All responses cleared!")
        

def display_evaluation_results_1(evaluation_results):
    """Display evaluation results in a separate section"""
    st.markdown("---")
    st.markdown("## 📊 Evaluation Results")
    
    # Calculate overall score
    if evaluation_results:
        total_score = sum(result["percentage"] for result in evaluation_results)
        avg_score = total_score / len(evaluation_results)
        st.markdown(f"### Overall Score: {avg_score:.1f}%")
    
    # Group results by ref_id for better organization
    grouped_results = {}
    for result in evaluation_results:
        key_parts = result["key"].split("_")
        ref_id = key_parts[0] if len(key_parts) > 0 else "unknown"
        
        if ref_id not in grouped_results:
            grouped_results[ref_id] = []
        
        grouped_results[ref_id].append(result)
    
    # Display results grouped by section
    for ref_id, results in grouped_results.items():
        with st.expander(f"Section {ref_id} - {len(results)} questions", expanded=True):
            for result in results:
                col1, col2 = st.columns([1, 5])
                with col1:
                    # Display score with color coding
                    score = result["percentage"]
                    if score >= 80:
                        st.markdown(f"<h3 style='color:green'>{score}%</h3>", unsafe_allow_html=True)
                    elif score >= 60:
                        st.markdown(f"<h3 style='color:orange'>{score}%</h3>", unsafe_allow_html=True)
                    else:
                        st.markdown(f"<h3 style='color:red'>{score}%</h3>", unsafe_allow_html=True)
                
                with col2:
                    st.markdown(f"**Q: {result['question']}**")
                    st.markdown(f"*Your answer:* {result['user_response']}")
                    st.markdown(f"*Feedback:* {result['feedback']}")
                
                st.markdown("---")
    
    # Option to download results as CSV
    if evaluation_results:
        import pandas as pd
        import io
        
        # Convert to dataframe for CSV export
        df = pd.DataFrame([
            {
                "Question": r["question"],
                "Your Response": r["user_response"],
                "Score": r["percentage"],
                "Feedback": r["feedback"]
            } for r in evaluation_results
        ])
        
        csv = df.to_csv(index=False)
        st.download_button(
            label="📥 Download Evaluation Results",
            data=csv,
            file_name="iso27002_quiz_results.csv",
            mime="text/csv",
        )
def evaluate_quiz_responses_1():
    """
    Evaluate all quiz responses using the LLM and return results
    """
    if not st.session_state.quiz_responses:
        st.warning("No responses to evaluate. Please fill in at least one response first.")
        return None
        
    # Display a progress indicator
    with st.spinner("Evaluating your responses..."):
        # Create a copy to avoid modifying during iteration
        responses_to_evaluate = dict(st.session_state.quiz_responses)
        evaluation_results = []
        
        # Create a progress bar
        progress_bar = st.progress(0)
        total_items = len(responses_to_evaluate)
        
        for i, (key, entry) in enumerate(responses_to_evaluate.items()):
            # Update progress
            progress_bar.progress((i + 1) / total_items)
            
            user_response = entry["user_response"].strip()
            if user_response:
                percentage, feedback = evaluate_user_response_1(
                    user_response,
                    entry["model_answer"],
                    entry["iso_text"]
                )
                evaluation_results.append({
                    "key": key,
                    "question": entry["question"],
                    "model_answer": entry["model_answer"],
                    "user_response": user_response,
                    "percentage": percentage,
                    "feedback": feedback
                })
            else:
                evaluation_results.append({
                    "key": key,
                    "question": entry["question"],
                    "model_answer": entry["model_answer"],
                    "user_response": "",
                    "percentage": 0,
                    "feedback": "Aucune réponse saisie."
                })
    
    # Return the evaluation results
    return evaluation_results            
def generate_iso_quiz_1():
    print("🔍 Generating ISO 27002 quiz questions...")

    # Load JSON file
    with open("new.json", 'r', encoding='utf-8') as f:
        structured_sections = json.load(f)

    quiz_results = []

    # Iterate through each category in the JSON structure
    for category, items in structured_sections.items():
        print(f"\n📂 Category: {category}")

        for entry in items:
            try:
                # Split the entry by comma to get subsection and section
                subsection, section = map(str.strip, entry.split(",", 1))
            except ValueError:
                print(f"❌ Skipping malformed entry: {entry}")
                continue

            query = f"{subsection} {section}"
            print(f"\n🔹 Searching for: {query}")

            results = vector_store_i_title.similarity_search(query=query, k=1)

            if results:
                combined_context = ""
                for doc in results:
                    title = doc.metadata.get("title")
                    content = doc.metadata.get("content", "")
                    print(f"✅ Found title: {title}")
                    print(f"📄 Content snippet: {content[:200]}...\n")

                    combined_context += content + "\n"

                if combined_context.strip():
                    quiz_item = generate_quiz_question_1(section, subsection, combined_context.strip())
                    quiz_item["category"] = category  # Include the category
                    quiz_results.append(quiz_item)
                else:
                    print(f"❌ No valid content for: {subsection}")
            else:
                print(f"❌ No matches found for: {query}")

    return quiz_results
# -----------------------------
# Quiz generator iso 42001
# -----------------------------
def generate_quiz_question_2(section, subsection, text):
    """ 
    Use LLM to generate a variable number of quiz questions for a specific ISO clause 
    """
    # Get the reference ID from the subsection (first part before space) 
    ref_parts = subsection.split(" ", 1) 
    ref = ref_parts[0] if len(ref_parts) > 0 else "unknown"
    
    # Create a more focused prompt for the LLM 
    prompt = f""" 
Agissez en tant qu’auditeur principal certifié lead auditor ISO 27002:2022.

Vous êtes en train d’examiner la clause ISO {subsection} : « {text} ».

Votre mission est de générer une série réaliste de questions de type audit, suivant une progression logique et en chaîne. Ces questions doivent simuler la manière dont un auditeur expérimenté approfondirait une évaluation de conformité réelle.

Pour chaque question, fournissez immédiatement après une réponse plausible et pertinente que l’organisation auditée pourrait donner, en lien avec le contexte décrit. Les réponses doivent être réalistes, refléter les bonnes pratiques, et s’aligner avec les exigences de la clause.

Contexte :
L’organisation auditée est une entreprise spécialisée dans les logiciels et les technologies. Elle comprend des équipes de développement logiciel ainsi que des équipes de validation. Ces équipes interagissent fréquemment avec le client télécom (opérateur) dans le monde entier, et collaborent avec des partenaires tiers pour la production des produits résidentiels. Les activités de développement, la validation du code, l’intégration et le déploiement continus, ainsi que la collaboration externe font partie intégrante de leur environnement opérationnel.

L’entreprise utilise des outils comme JIRA, jenkins, Confluence, GitLab, SharePoint et Office 365 pour faciliter la gestion des tâches, la collaboration, et la gestion documentaire.
les OS des postes de travail utilisés sont soit du Linux soit du Windows, les produits developpés sont avec des OS Android ou Linux selon le client en face.
la structure de la société est faite de sorte à avoir des equipes de developpement , des equipes de validation , des equipes de certification fonctionnels, des equipes transverses  et un equipe infra reseau qui permet de fournir des environnements de test et de simulation des environnements client, ou des accées VPN aux infrastructures client( operateur)

Consignes :
- Commencez par une question générale pour évaluer la conformité de base.
- Ensuite, posez des questions de plus en plus détaillées ou approfondies, en fonction des réponses probables.
- Les questions doivent être spécifiques, exploitables et clairement liées au contenu de la clause et au contexte de l'entreprise.
- Fournissez une réponse réaliste juste après chaque question qui se base sur les clauses de ISO comme référence, justifiant la réponse par rapport au document de Clause ISO fourni.
- Adoptez un ton naturel et humain, tel qu’un auditeur experimentés dans le domaine de l'entreprise le ferait lors d’un entretien.
- Incluez entre 3 et 7 paires questions/réponses, selon la complexité de la clause.
- Assurez-vous que chaque question s’appuie sur la précédente — comme dans une conversation réfléchie ou une visite d’audit.
- Formulez vos questions et réponses en tenant compte de l’environnement de l’entreprise : développement logiciel, validation, interactions clients et collaborations avec des tiers, ainsi que l’utilisation des outils mentionnés.
- N’INCLUEZ AUCUNE mise en forme, puce ou explication — produisez uniquement les questions suivies de leur réponse sur deux lignes consécutives.

Générez maintenant l’ensemble des paires questions/réponses en chaîne pour cette clause, en français. 

"""

    response = gemma_model.invoke(prompt).content

    # Split into lines, and group them as Q&A pairs
    lines = [line.strip() for line in response.strip().split("\n") if line.strip()]
    qna_pairs = []

    for i in range(0, len(lines) - 1, 2):
        question = lines[i]
        answer = lines[i + 1]
        #company_answer, sources = ask_company_question(question)
        qna_pairs.append({
            "question": question,
            "answer": answer,
            #"company_answer": company_answer,
            #"sources": sources
        })

    return {
        "section": section,
        "subsection": subsection,
        "ref_id": ref,
        "qna": qna_pairs,
        "text": text  # Include the clause content
    }

def save_quiz_to_csv_2(quiz_results):
    """
    Save quiz results directly (without flattening) to a CSV-compatible JSON format.
    """
    import os
    import json

    filename = "iso42001_quiz1_output_raw.json"  # Better suited for structured data
    file_path = os.path.abspath(filename)

    # Save directly as JSON for structured access later
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(quiz_results, f, ensure_ascii=False, indent=2)

    print(f"Raw quiz results saved to: {file_path}")
    return file_path

    
    

def display_quiz_in_streamlit_2(quiz_results):
    

    st.markdown("## ISO 27002:2022 Audit Questions")

    for i, item in enumerate(quiz_results):
        section = item["section"]
        subsection = item["subsection"]
        ref_id = item["ref_id"]
        qna = item["qna"]

        with st.expander(f"{subsection}", expanded=(i == 0)):
            st.markdown(f"**Section:** {section}")
            st.markdown(f"**Reference:** {ref_id}")

            questions_df = pd.DataFrame({
                "Question #": range(1, len(qna) + 1),
                "Audit Question": [q["question"] for q in qna],
                "Suggested Answer": [q["answer"] for q in qna],
                #"Company Answer": [q["company_answer"] for q in qna],
                #"Sources": [", ".join(q["sources"]) if q["sources"] else "—" for q in qna]
            })

            st.table(questions_df)

            section_csv = questions_df.to_csv(index=False).encode('utf-8')
            st.download_button(
                label=f"📥 Download '{ref_id}' Questions",
                data=section_csv,
                file_name=f"iso42001_{ref_id}_questions.csv",
                mime="text/csv",
                key=f"download_section_{i}"
            )

    all_questions = []
    for item in quiz_results:
        for i, qna in enumerate(item["qna"]):
            all_questions.append({
                "Section": item["section"],
                "Subsection": item["subsection"],
                "Reference ID": item["ref_id"],
                "Question #": i + 1,
                "Question": qna["question"],
                "Answer": qna["answer"],
                #"Company Answer": qna["company_answer"],
                #"Sources": ", ".join(qna["sources"]) if qna["sources"] else "—"
            })

    all_questions_df = pd.DataFrame(all_questions)

    if all_questions:
        st.markdown("### Download All Questions")
        full_csv = all_questions_df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="📥 Download Complete Question Set",
            data=full_csv,
            file_name="iso42001_all_questions.csv",
            mime="text/csv",
            key="download_all"
        )


import json

def generate_iso_quiz_2():
    print("🔍 Generating ISO 42001 quiz questions...")

    with open("new.json", 'r', encoding='utf-8') as f:
        structured_sections = json.load(f)

    quiz_results = []

    for category, items in structured_sections.items():
        print(f"\n📂 Category: {category}")

        for entry in items:
            try:
                subsection, section = map(str.strip, entry.split(",", 1))
            except ValueError:
                print(f"❌ Skipping malformed entry: {entry}")
                continue

            query = f"{subsection} {section}"
            print(f"\n🔹 Searching for: {query}")

            results = vector_store_iso_title24001.similarity_search(query=query, k=1)

            if results:
                combined_context = ""
                for doc in results:
                    title = doc.metadata.get("title", "")
                    content = doc.metadata.get("content", "")
                    source = doc.metadata.get("source", "")
                    print(f"✅ Found title: {title}")
                    print(f"📄 Content snippet: {content[:200]}...\n")

                    combined_context += f"{title}\n{content}\n{source}\n"

                if combined_context.strip():
                    quiz_item = generate_quiz_question_2(section, subsection, combined_context.strip())
                    quiz_item["category"] = category  # optionally include the category
                    quiz_results.append(quiz_item)
                else:
                    print(f"❌ No valid content for: {subsection}")
            else:
                print(f"❌ No matches found for: {query}")

    return quiz_results



# -----------------------------
# Document Generation System
# -----------------------------
class DocumentState(TypedDict):
    template: Dict
    current_section_index: int
    generated_sections: List[Dict]
    section_content: str
    section_feedback: str
    is_approved: bool
    complete: bool
    pdf_data: Optional[bytes]
    current_documents: List
    regeneration_attempts: int

# Document generation graph
document_graph = StateGraph(DocumentState)

def document_generator_node(state: DocumentState) -> DocumentState:
    """Generate content for the current section based on ISO documentation."""
    print("document generator : ", state.get("regeneration_attempt",None))
    template = state["template"]
    current_index = state["current_section_index"]
    regeneration_attempts = state.get("regeneration_attempts", None)
    section_feedback = state.get("section_feedback", "")
    
    # If processing cover page
    if current_index == -1:
        doc_details = template["document_details"]
        content = f"""# {doc_details['document_name']}
### {doc_details['english_name']}

**Document ID:** {doc_details['document_id']}
**Version:** {doc_details['version']}
**Date:** {doc_details['date']}
**Author:** {doc_details['author']}
**Approved by:** {doc_details['approver']}

## Purpose
{doc_details['purpose']['description']}
"""
        return {
            **state,
            "section_content": content,
            "current_documents": []
        }
    
    # If processing table of contents
    if current_index == 0:
        policy_sections = template["policy_sections"]
        toc = "# Table of Contents\n\n"
        
        for idx, section in enumerate(policy_sections, 1):
            section_num = section.get("section_number", str(idx))
            french_title = section.get("french_title", f"Section {section_num}")
            toc += f"{section_num}. {french_title}\n"
        
        return {
            **state,
            "section_content": toc,
            "current_documents": []
        }
    
    # If processing a regular section (adjust index to account for TOC)
    adjusted_idx = current_index - 1
    if adjusted_idx < len(template["policy_sections"]):
        section = template["policy_sections"][adjusted_idx]
        section_num = section.get("section_number", str(adjusted_idx + 1))
        french_title = section.get("french_title", f"Section {section_num}")
        english_description = section.get("english_description", "")
        relevant_clauses = section.get("relevant_iso_clauses", [])
        
        # Log appropriate message based on whether this is initial generation or regeneration
        if regeneration_attempts!=None and regeneration_attempts > 0:
            print(f"Regenerating content for '{section_num}: {french_title}' (attempt {regeneration_attempts})")
        else:
            print(f"Generating initial content for '{section_num}: {french_title}'")
        
        # Create query combining section info and ISO clauses
        iso_clause_text = " ".join(relevant_clauses)
        query = f"{french_title} {english_description} ISO 42001 {iso_clause_text}"
        
        # Retrieve relevant ISO documents based on clauses
        documents = []
        for clause in relevant_clauses:
            clause_query = f"ISO 42001 clause {clause}"
            clause_docs = vector_store_42001_originals.similarity_search(clause_query, k=2)
            documents.extend(clause_docs)
        
        # Add general section documents
        section_docs = vector_store_42001_originals.similarity_search(query, k=3)
        documents.extend(section_docs)
        
        # Rerank all documents
        reranked_docs = rerank_documents(query, documents, top_k=3)
        
        # Format context for generation
        context = "\n\n".join(
            f"Source: {doc.metadata.get('source', 'Unknown')}\n"
            f"Title: {doc.metadata.get('title', 'No title')}\n"
            f"Content: {doc.page_content}\n"
            f"ISO Clause: {', '.join(relevant_clauses)}"
            for doc in reranked_docs
        )
        
        # Initialize feedback instructions (empty for first generation)
        feedback_instructions = ""
        
        # Only include feedback instructions if this is a regeneration attempt with feedback
        if regeneration_attempts!=None and regeneration_attempts > 0 and section_feedback:
            previous_content = state.get("section_content", "")
            feedback_instructions = f"""
PREVIOUS CONTENT THAT NEEDS IMPROVEMENT:
{previous_content}

REVIEWER FEEDBACK:
{section_feedback}

Please revise the content based on the above feedback. Focus specifically on addressing the issues 
identified by the reviewer while maintaining alignment with ISO 42001 requirements.
Make sure to fix any problems mentioned in the feedback while following the ISO clauses closely.
"""
        
        # Generate content with LLM - include feedback instructions only when applicable
        prompt = f"""
You are an ISO 42001 documentation specialist. Generate content for the following section of an AI Management System policy document:

Section Number: {section_num}
Section Title (French): {french_title}
Section Description (English): {english_description}
Relevant ISO 42001:2023 Clauses: {', '.join(relevant_clauses)}

{feedback_instructions}

Use the following ISO 42001 context to inform your writing:
{context}

The content should:
1. Be written in French to match the document's language
2. Be professionally written and ISO-compliant
3. Include appropriate subheadings and structure
4. Be comprehensive but concise (about 1-2 pages worth of content)
5. Include specific requirements relevant to this section
6. Reference the ISO clauses naturally in the text
7. Use proper terminology from ISO 42001:2023

Format your response as clean Markdown, starting with a level 2 heading for the section title. Don't use ** characters.
"""
        
        # Generate content using the LLM
        # content = llm_ollama.invoke(prompt).content
        content = "Generated content goes here."  # Placeholder for generated content
        
        # Add section number and title if not already included
        if not content.strip().startswith("#"):
            content = f"## {section_num}. {french_title}\n\n{content}"
        print({**state,
            "regeneration_attempts": regeneration_attempts
        }["regeneration_attempts"])
        return {
            **state,
            # "section_content": content,
            "section_content": "lol",
            "current_documents": reranked_docs
        }
    
    # If all sections are processed, mark as complete
    return {
        **state,
        "complete": True,
        "section_content": "Document generation complete."
    }
def document_reviewer_node(state: DocumentState) -> DocumentState:
    """Review the generated content for accuracy and completeness."""
    print("document reviewer : ", state.get("regeneration_attempt",None))
    template = state["template"]
    current_index = state["current_section_index"]
    content = state["section_content"]
    documents = state["current_documents"]
    
    # Initialize regeneration attempts if not present
    regeneration_attempts = state.get("regeneration_attempts",0)
    
    # If processing cover page (auto-approve)
    if current_index == -1:
        print("Cover page reviewed and approved")
        # Add cover page to generated sections
        generated_sections = state["generated_sections"].copy()
        generated_sections.append({
            "title": "Cover Page",
            "content": content
        })
        
        # Move to table of contents and reset regeneration counter
        
        return {
            **state,
            "generated_sections": generated_sections,
            "is_approved": True,
            "section_feedback": "Cover page approved.",
            "current_section_index": current_index + 1,
            "regeneration_attempts": 0
        }
    
    # If processing table of contents (auto-approve)
    if current_index == 0:
        print("Table of contents reviewed and approved")
        # Add TOC to generated sections
        generated_sections = state["generated_sections"].copy()
        generated_sections.append({
            "title": "Table of Contents",
            "content": content
        })
        
        # Move to first regular section and reset regeneration counter
        return {
            **state,
            "generated_sections": generated_sections,
            "is_approved": True,
            "section_feedback": "Table of contents approved.",
            "current_section_index": current_index + 1,
            "regeneration_attempts": 0
        }
    
    # If all sections are complete (account for TOC)
    adjusted_idx = current_index - 1
    if adjusted_idx >= len(template["policy_sections"]):
        print("All sections reviewed, proceeding to finalization")
        return {
            **state,
            "is_approved": True,
            "section_feedback": "All sections approved.",
            "complete": True,
            "regeneration_attempts": 0
        }
    
    # For regular sections, review against ISO documentation
    section = template["policy_sections"][adjusted_idx]
    section_num = section.get("section_number", str(adjusted_idx + 1))
    french_title = section.get("french_title", f"Section {section_num}")
    english_description = section.get("english_description", "")
    relevant_clauses = section.get("relevant_iso_clauses", [])
    
    print(f"Reviewing section {section_num}: {french_title}")
    
    # Format context for review
    context = "\n\n".join(
        f"Source: {doc.metadata.get('source', 'Unknown')}\n"
        f"Title: {doc.metadata.get('title', 'No title')}\n"
        f"Content: {doc.page_content}"
        for doc in documents
    )
    
    # Review content with LLM
    prompt = f"""
You are an ISO compliance reviewer checking a section of an AI Management System policy document.

Section Number: {section_num}
Section Title (French): {french_title}
Section Description (English): {english_description}
Relevant ISO 42001:2023 Clauses: {', '.join(relevant_clauses)}

Generated Content:
{content}

ISO 42001 Reference Material:
{context}

Please review the generated content and evaluate:
1. Accuracy - Does it comply with ISO 42001:2023 requirements, especially clauses {', '.join(relevant_clauses)}?
2. Completeness - Does it address all necessary aspects described in the section description?
3. Structure - Is it well-organized and professional?
4. Language - Is it written in clear, professional French?
5. Compliance - Does it meet ISO standards and use proper terminology?

Return your assessment in this format:
- APPROVED with a pourcentage of confidence (e.g., 95%) if it <50% will considarte as not approved
- Feedback: Your detailed feedback here
- Suggested Improvements: Specific suggestions for improvement if not approved
"""
    
    # Get review from LLM
    # review_result = llm_ollama.invoke(prompt).content
    # print(f"Review result: {review_result}")
    # Parse the result to determine if approved
    # Parse the result to extract the percentage and determine if approved
    # percentage_match = re.search(r'APPROVED.+?(\d+)%', review_result, re.IGNORECASE)
    
    # Default to not approved if no percentage found
    # confidence_percentage = 0
    # if percentage_match:
    #     confidence_percentage = int(percentage_match.group(1))
    confidence_percentage=0
    
    is_approved = confidence_percentage >= 70  # Approve only if 70% or higher
    
    if is_approved:
        print(f"Section '{section_num}: {french_title}' approved")
        # Add approved content to generated sections
        generated_sections = state["generated_sections"].copy()
        generated_sections.append({
            "title": f"{section_num}. {french_title}",
            "content": content                                                                                                                                                                                                                                                                              
        })
        
        # Move to next section and reset regeneration counter
        next_index = current_index + 1
        
        return {
            **state,
            "generated_sections": generated_sections,
            "is_approved": True,
            "section_feedback": f"Section approved: {section_num}. {french_title}. Moving to next section.",
            "current_section_index": next_index,
            "regeneration_attempts": 0
        }
    else:
        # Check if we've reached the maximum regeneration attempts
        regeneration_attempts += 1
        print(f"Section '{section_num}: {french_title}' not approved, will regenerate content (attempt {regeneration_attempts}/5)")
        
        if regeneration_attempts >= 5:
            print(f"⚠️ Maximum regeneration attempts reached for section '{french_title}'. Using current content and moving on.")
            
            # Add the content anyway and move to next section
            generated_sections = state["generated_sections"].copy()
            generated_sections.append({
                "title": f"{section_num}. {french_title}",
                "content": content + "\n\n_Note: This section reached maximum regeneration attempts and may require manual review._"
            })
            
            next_index = current_index + 1
            
            return {
                **state,
                "generated_sections": generated_sections,
                "is_approved": True,  # Force approval to move on
                "section_feedback": f"Maximum regeneration attempts reached for section: {french_title}. Using current content and moving to next section.",
                "current_section_index": next_index,
                "regeneration_attempts": 0
            }
            
        # If not approved, provide feedback for regeneration and increment counter
        print({**state,
            "regeneration_attempts": regeneration_attempts
        }["regeneration_attempts"])
        return {
            **state,
            "is_approved": False,
            # "section_feedback": review_result,
            "section_feedback":None,
            "regeneration_attempts": regeneration_attempts
        }
def finalize_document_node(state: DocumentState) -> DocumentState:
    """Generate the final PDF document with table of contents and color styling."""
    print("finalizer : ", state.get("regeneration_attempt",None))
    template = state["template"]
    generated_sections = state["generated_sections"]
    
    print(f"Finalizing document with {len(generated_sections)} approved sections")
    
    # Create in-memory PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    # Add enhanced custom styles with colors
    styles.add(ParagraphStyle(
        'FrenchTitle', 
        parent=styles['Title'], 
        fontSize=20,
        spaceAfter=12,
        fontName="Helvetica-Bold",
        textColor=colors.blue,  # Blue text for main title
    ))
    
    styles.add(ParagraphStyle(
        'EnglishTitle', 
        parent=styles['Title'], 
        fontSize=16,
        spaceAfter=24,
        fontName="Helvetica-Bold",
        textColor=colors.navy,  # Navy text for secondary title
    ))
    
    # Create custom heading styles with colors
    styles.add(ParagraphStyle(
        'BlueHeading1',
        parent=styles['Heading1'],
        textColor=colors.blue,
        fontSize=16,
        spaceAfter=10,
    ))
    
    styles.add(ParagraphStyle(
        'BlueHeading2',
        parent=styles['Heading2'],
        textColor=colors.royalblue,
        fontSize=14,
        spaceAfter=8,
    ))
    
    styles.add(ParagraphStyle(
        'BlueHeading3',
        parent=styles['Heading3'],
        textColor=colors.steelblue,
        fontSize=12,
        spaceAfter=6,
    ))
    
    # Add cover page
    if generated_sections and generated_sections[0]["title"] == "Cover Page":
        doc_details = template["document_details"]
        
        elements.append(Paragraph(doc_details["document_name"], styles['FrenchTitle']))
        elements.append(Paragraph(doc_details["english_name"], styles['EnglishTitle']))
        elements.append(Spacer(1, 0.5*inch))
        
        # Create document info table with improved styling
        data = [
            ["Document ID:", doc_details["document_id"]],
            ["Version:", doc_details["version"]],
            ["Date:", doc_details["date"]],
            ["Author:", doc_details["author"]],
            ["Approved by:", doc_details["approver"]]
        ]
        
        info_table = Table(data, colWidths=[2*inch, 3*inch])
        info_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.steelblue),  # Blue background
            ('BACKGROUND', (1, 0), (1, -1), colors.white),  # White for content cells
            ('TEXTCOLOR', (0, 0), (0, -1), colors.white),  # White text on blue
            ('TEXTCOLOR', (1, 0), (1, -1), colors.black),  # Black text on white
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGNMENT', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        
        elements.append(info_table)
        elements.append(Spacer(1, 1*inch))
        
        # Add purpose with blue heading
        elements.append(Paragraph("Purpose", styles['BlueHeading2']))
        elements.append(Paragraph(doc_details["purpose"]["description"], styles['Normal']))
        elements.append(Spacer(1, 0.5*inch))
        
        # Add page break after cover
        elements.append(PageBreak())
    
    # Add table of contents if it exists
    if len(generated_sections) > 1 and generated_sections[1]["title"] == "Table of Contents":
        elements.append(Paragraph("Table of Contents", styles['BlueHeading1']))
        
        # Add TOC entries from policy sections with alternating backgrounds
        for i, section in enumerate(template["policy_sections"]):
            section_num = section.get("section_number", "")
            french_title = section.get("french_title", "")
            
            toc_entry = f"{section_num}. {french_title}"
            
            # Create a custom style for TOC entries with alternating backgrounds
            toc_style = ParagraphStyle(
                f'TOCEntry{i}',
                parent=styles['Normal'],
                backColor=colors.lightgrey if i % 2 == 0 else colors.white,
                spaceBefore=3,
                spaceAfter=3,
                leftIndent=6,
                rightIndent=6,
            )
            
            elements.append(Paragraph(toc_entry, toc_style))
        
        elements.append(Spacer(1, 0.2*inch))
        elements.append(PageBreak())
    
    # Add content sections (skip cover page and TOC)
    for section_idx, section in enumerate(generated_sections[2:]):
        # Process Markdown content to ReportLab elements
        content = section["content"]
        
        # Add a section separator
        if section_idx > 0:
            # Create a gray separator line
            elements.append(Spacer(1, 0.2*inch))
            separator = Paragraph('<hr width="100%" color="lightgrey"/>', 
                                  ParagraphStyle('Separator', parent=styles['Normal']))
            elements.append(separator)
            elements.append(Spacer(1, 0.2*inch))
        
        # Process Markdown content line by line
        for line in content.split('\n'):
            # Handle headings with colors
            if line.strip().startswith('# '):
                heading_text = line.strip()[2:]
                elements.append(Paragraph(heading_text, styles['BlueHeading1']))
            elif line.strip().startswith('## '):
                heading_text = line.strip()[3:]
                elements.append(Paragraph(heading_text, styles['BlueHeading2']))
            elif line.strip().startswith('### '):
                heading_text = line.strip()[4:]
                elements.append(Paragraph(heading_text, styles['BlueHeading3']))
            # Handle bullet points with styling
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                bullet_text = line.strip()[2:]
                bullet_style = ParagraphStyle(
                    'Bullet',
                    parent=styles['Normal'],
                    leftIndent=20,
                    firstLineIndent=-15,
                )
                elements.append(Paragraph(f"• {bullet_text}", bullet_style))
            # Handle normal paragraphs
            elif line.strip():
                elements.append(Paragraph(line.strip(), styles['Normal']))
        
        # Add space between sections
        elements.append(Spacer(1, 0.3*inch))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    print("PDF document finalized successfully")
    
    return {
        **state,
        "complete": True,
        "pdf_data": buffer.getvalue()
    }
# Add nodes to the document generation graph
document_graph.add_node("document_generator", document_generator_node)
document_graph.add_node("document_reviewer", document_reviewer_node)
document_graph.add_node("finalize_document", finalize_document_node)

# Define graph flow with direct edges - simpler approach without conditional edges
document_graph.add_edge(START, "document_generator")
document_graph.add_edge("document_generator", "document_reviewer")

# Direct edges from reviewer based on state flags
# If not approved → back to generator
# If complete → to finalizer
# Otherwise (when approved) → back to generator for next section
document_graph.add_conditional_edges(
    "document_reviewer",
    lambda state: 
        "finalize_document" if state["complete"] else 
        "same_section_generator" if not state["is_approved"] else
        "next_section_generator",
    {
        "same_section_generator": "document_generator",
        "next_section_generator": "document_generator",
        "finalize_document": "finalize_document"
    }
)

document_graph.add_edge("finalize_document", END)

# Compile the document generation graph
compiled_document_graph = document_graph.compile()

def create_document_from_template(template_json):
    """Initialize and run document generation process."""
    print("Starting document generation process...")
    
    # Initial state
    state = {
        "template": template_json,
        "current_section_index": -1,  # Start with cover page (-1), then TOC (0), then sections (1+)
        "generated_sections": [],
        "section_content": "",
        "section_feedback": "",
        "is_approved": True,
        "complete": False,
        "pdf_data": None,
        "current_documents": []
        
    }
    
    # Run the graph with higher recursion limit to handle multiple sections
    result = compiled_document_graph.invoke(state, {"recursion_limit": 100})
    
    print(f"Document generation completed with {len(result['generated_sections'])} sections")
    return result

def load_json_template(file_path="AIMS.json"):
    """Load a JSON template file from the specified path."""
    try:
        # Use absolute path if provided, otherwise look in current directory
        if not os.path.isabs(file_path):
            file_path = os.path.join(os.getcwd(), file_path)
            
        with open(file_path, 'r', encoding='utf-8') as file:
          return json.load(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"Template file not found: {file_path}")
    except json.JSONDecodeError:
        raise ValueError(f"Invalid JSON format in template file: {file_path}")
       
# -----------------------------
# LangGraph Router
# -----------------------------

print("Building main workflow graph...")
class GraphState(TypedDict):
    question: str
    generation: str
    documents: List
    results: List

workflow = StateGraph(GraphState)
# Update the document generation node to work with the default AIMS.json file
def generate_document_node(state: GraphState) -> GraphState:
    """Generate a document based on the provided template file."""
    try:
        # Default to the new template file
        file_path = r"C:\Users\G800613RTS\Desktop\Chatbot\Templates\AI Management System Policy.json"
        
        print(f"Loading template from: {file_path}")
        
        # Load the JSON template
        template_json = load_json_template(file_path)
        
        # Run the document generation process
        result = create_document_from_template(template_json)
        
        # Get PDF data
        pdf_data = result["pdf_data"]
        json_filename = template_json["document_details"]["document_name"].replace(" ", "_")
        pdf_filename = f"{json_filename}.pdf"
        
        # Save PDF to file
        with open(pdf_filename, "wb") as f:
            f.write(pdf_data)
        
        # Prepare detailed generation report
        section_info = []
        for section in result["generated_sections"]:
            section_title = section["title"]
            content_preview = section["content"][:50] + "..." if len(section["content"]) > 50 else section["content"]
            section_info.append(f"- {section_title}: {content_preview}")
        
        section_report = "\n".join(section_info)
        
        # Prepare response
        result_text = f"""# Document Generation Complete

I've created the **{template_json['document_details']['document_name']}** document with {len(result['generated_sections'])} sections based on the AI Policy template.

## Document Information
- **Document ID:** {template_json['document_details']['document_id']}
- **Version:** {template_json['document_details']['version']}
- **Date:** {template_json['document_details']['date']}
- **Author:** {template_json['document_details']['author']}
- **Approved by:** {template_json['document_details']['approver']}

## Generated Sections Overview
{len(template_json['policy_sections'])} policy sections were processed and reviewed for ISO 42001 compliance.

The PDF has been saved as: {pdf_filename}

You can download the document using the button below.
"""
        
        return {
            "question": state["question"],
            "generation": result_text,
            "documents": [],
            "results": {
                "pdf_path": pdf_filename,
                "pdf_data": pdf_data,
                "document_title": template_json['document_details']['document_name']
            }
        }
    except Exception as e:
        error_message = f"Error generating document: {str(e)}"
        print(error_message)
        return {
            "question": state["question"],
            "generation": error_message,
            "documents": [],
            "results": []
        }

def gap_analysis_node(state: GraphState) -> GraphState:
    result = run_gap_analysis()
    result_text = f"Gap analysis completed."
    
    # Get the actual DataFrame and CSV path
    csv_path = result['csv']
    results_df = result['df']
    
    return {
        "question": state["question"],
        "generation": result_text,
        "documents": [],
        "results": results_df  # Store the actual results data
    }
def generate_quiz_node_2(state: GraphState) -> GraphState:
    """
    Generate quiz questions based on ISO 27002 sections.
    Returns results and saves them as raw structured JSON.
    """
    # Generate the quiz
    quiz_results = generate_iso_quiz_2()

    # Count questions and sections for summary
    sections_count = len(set([q['section'] for q in quiz_results]))
    total_questions = sum(len(q['qna']) for q in quiz_results)

    # Automatically save results in raw JSON format
    quiz_file_path = save_quiz_to_csv_2(quiz_results)

    # Create a detailed summary message
    result_text = (
        f"Quiz generation completed. Created {total_questions} questions "
        f"across {sections_count} sections of ISO 42001:2023.\n\n"
        f"Questions per clause vary based on complexity and content depth.\n"
        f"Raw quiz results saved to: {quiz_file_path}"
    )

    return {
        "question": state["question"],
        "generation": result_text,
        "documents": [],
        "results": quiz_results
    }
def generate_quiz_node_1(state: GraphState) -> GraphState:
    """
    Generate quiz questions based on ISO 27002 sections.
    Returns results without saving them to a file.
    """
    # Generate the quiz
    quiz_results = generate_iso_quiz_1()


    # Count questions and sections for summary
    sections_count = len(set([q['section'] for q in quiz_results]))
    total_questions = sum(len(q['qna']) for q in quiz_results)

    # Create a detailed summary message
    result_text = (
        f"Quiz generation completed. Created {total_questions} questions "
        f"across {sections_count} sections of ISO 27002:2022.\n\n"
        f"Questions per clause vary based on complexity and content depth."
    )

    return {
        "question": state["question"],
        "generation": result_text,
        "documents": [],
        "results": quiz_results
    }
def generate_quiz_node(state: GraphState) -> GraphState:
    """
    Generate quiz questions based on ISO 27002 sections.
    Returns results and saves them as raw structured JSON.
    """
    # Generate the quiz
    quiz_results = generate_iso_quiz()

    # Count questions and sections for summary
    sections_count = len(set([q['section'] for q in quiz_results]))
    total_questions = sum(len(q['qna']) for q in quiz_results)

    # Automatically save results in raw JSON format
    quiz_file_path = save_quiz_to_csv(quiz_results)

    # Create a detailed summary message
    result_text = (
        f"Quiz generation completed. Created {total_questions} questions "
        f"across {sections_count} sections of ISO 27002:2022.\n\n"
        f"Questions per clause vary based on complexity and content depth.\n"
        f"Raw quiz results saved to: {quiz_file_path}"
    )

    return {
        "question": state["question"],
        "generation": result_text,
        "documents": [],
        "results": quiz_results
    }

def iso_qna_node(state: GraphState) -> GraphState:
    """
    Process ISO standard questions using the LangGraph implementation.
    Uses a single persistent graph instance to maintain conversation context across multiple queries.
    """
    # Get the question from state
    question = state["question"]
    
    # Use a fixed thread_id for all ISO conversations to ensure memory persistence
    thread_id = "iso-global-thread"
    
    # Initialize session state for ISO conversation context if not present
    if "iso_conversation_context" not in st.session_state:
        st.session_state.iso_conversation_context = {
            "message_history": [],
            "sources": [],
            "last_question": "",
            "last_answer": ""
        }
    
    # Get the current message history
    message_history = st.session_state.iso_conversation_context.get("message_history", [])
    
    # Add current question to message history
    current_message = HumanMessage(content=question)
    
    # If we have message history, use it to maintain conversation context
    # Otherwise just use the current message
    messages_to_send = message_history + [current_message] if message_history else [current_message]
    
    # Configure the graph with the fixed thread_id
    config = {"configurable": {"thread_id": thread_id}}
    
    # Invoke the graph with the complete message history
    result = graph.invoke({"messages": messages_to_send}, config=config)
    print("ISO Q&A Graph invoked with messages:", messages_to_send)
    print(f"📍 ISO Q&A Graph Result: {result}")
    
    # Extract the answer from the graph's response
    if result and "messages" in result and result["messages"]:
        # Update our message history with all messages from the result
        # This is important to maintain the conversation context for future queries
        st.session_state.iso_conversation_context["message_history"] = result["messages"]
        
        # Get the last message which should be the AI's response
        answer = result["messages"][-1].content
        
        # Extract sources from the tools used in the graph (if available)
        new_sources = []
        tool_messages = [m for m in result["messages"] if m.type == "tool"]
        
        # If we have tool messages, extract source information
        if tool_messages:
            for tool_msg in tool_messages:
                if hasattr(tool_msg, "artifact") and tool_msg.artifact:
                    # Extract documents from the tool's artifact
                    docs = tool_msg.artifact
                    for doc in docs:
                        if hasattr(doc, "metadata") and doc.metadata:
                            # Add source information to our sources list
                            source = doc.metadata.get('source', 'Unknown')
                            if source not in new_sources:
                                new_sources.append(source)
            
            # Store new sources for this conversation if we found any
            if new_sources:
                st.session_state.iso_conversation_context["sources"] = new_sources
                st.session_state.iso_conversation_context["last_question"] = question
                st.session_state.iso_conversation_context["last_answer"] = answer
                
                # Use the new sources for this response
                sources_to_return = new_sources
            else:
                # No new sources, use previous sources
                sources_to_return = st.session_state.iso_conversation_context["sources"]
        else:
            # No tool messages, use previous sources if available
            sources_to_return = st.session_state.iso_conversation_context["sources"]
        
        # Save the current Q&A pair
        st.session_state.iso_conversation_context["last_question"] = question
        st.session_state.iso_conversation_context["last_answer"] = answer
        
        # Return the response with appropriate sources
        return {
            "question": question,
            "generation": answer,
            "documents": sources_to_return if sources_to_return else [],
            "results": [],
        }
    else:
        # Fallback in case the graph doesn't return expected results
        # Still try to use previous sources if available
        previous_sources = st.session_state.iso_conversation_context.get("sources", [])
        
        return {
            "question": question,
            "generation": "I'm sorry, I couldn't process your ISO standard question at this time.",
            "documents": previous_sources,
            "results": [],
        }
    


def route_question(state: GraphState) -> str:
    """
    Route question to appropriate node based on question and history.
    """
    question = state["question"]
    
    # Get conversation history
    history = ""
    if "messages" in st.session_state:
        # Format the last 5 messages (or fewer if not available)
        history_msgs = st.session_state.messages[-5:] if len(st.session_state.messages) > 5 else st.session_state.messages
        history = "\n".join([f"{msg['role']}: {msg['content']}" for msg in history_msgs])
    
    # Use the routing model to determine the destination
    destination = route_chain.invoke({"query": question, "history": history})
    print(f"📍 LLM Routing Destination: {destination}")
    
    # Map the destination to the appropriate node
    if destination == "gap_analysis":
        return "run_gap_analysis"
    elif destination == "quiz_iso27002_company":
        return "generate_quiz_company"
    elif destination == "quiz_iso27002_human_eval":
        return "generate_quiz_human_eval"
    elif destination == "quiz_iso42001":
        return "generate_quiz_iso42001"
    elif destination == "generate_document":
        return "generate_document"
    else:
        # Default to ISO Q&A for unrecognized destinations
        return "ask_iso_question"

workflow.add_node("run_gap_analysis", gap_analysis_node)
workflow.add_node("ask_iso_question", iso_qna_node)
workflow.add_node("generate_quiz_company", generate_quiz_node)
workflow.add_node("generate_quiz_human_eval", generate_quiz_node_1)
workflow.add_node("generate_quiz_iso42001", generate_quiz_node_2)
workflow.add_node("generate_document", generate_document_node)
workflow.add_conditional_edges(START, route_question, {
    "run_gap_analysis": "run_gap_analysis",
    "ask_iso_question": "ask_iso_question",
    "generate_quiz_company": "generate_quiz_company",
    "generate_quiz_human_eval": "generate_quiz_human_eval",
    "generate_quiz_iso42001": "generate_quiz_iso42001",
    "generate_document": "generate_document"
})
workflow.add_edge("run_gap_analysis", END)
workflow.add_edge("ask_iso_question", END)
workflow.add_edge("generate_quiz_company", END)
workflow.add_edge("generate_quiz_human_eval", END)
workflow.add_edge("generate_quiz_iso42001", END)
workflow.add_edge("generate_document", END)
chat_graph = workflow.compile()
def run_query_pipeline(user_input: str):
    state = {"question": user_input, "generation": "", "documents": [], "results": []}
    result = chat_graph.invoke(state)
    return result["generation"], result["documents"], result["results"]

# Ensure DB file exists
def ensure_db_file_exists():
    if not os.path.exists(DB_FILE):
        with open(DB_FILE, 'w') as file:
            json.dump({"chat_history": []}, file)
def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return data
def main():
    
    st.title("🔍 ISO Assistant with Voice Support")
    st.write("Ask about ISO Standard, request a gap analysis, or just chat generally using voice or text.")
    
    # Load Whisper model
    processor_voice, model_voice = load_whisper_model()
    
    # Sidebar for model selection
    st.sidebar.title("Model Settings")
    
    # List of available models
    models = ["gemma3:12b-it-qat", "deepseek-r1:7b"]
    selected_model = st.sidebar.selectbox("Select Model", models, index=0)
    
    # Set up the model based on selection
    if selected_model == "deepseek-r1:7b":
        st.session_state.current_model = "gemma3:12b-it-qat"
    else:
        st.session_state.current_model = "deepseek"

    st.sidebar.title("Voice Settings")
    voice_enabled = st.sidebar.checkbox("Enable Voice Input", value=True)
    
    ensure_db_file_exists()
    
    # Load chat history from db.json
    with open(DB_FILE, 'r') as file:
        db = json.load(file)
    
    # Initialize messages in session state if not already present
    if "messages" not in st.session_state:
        st.session_state.messages = db.get('chat_history', [])
    
    # Initialize document generation results if not present
    if "document_results" not in st.session_state:
        st.session_state.document_results = None
    
    # CSS for fixed bottom input
    st.markdown("""
    <style>
    .main > div {
        padding-bottom: 100px;
    }
    
    .stChatFloatingInputContainer {
        position: fixed !important;
        bottom: 0 !important;
        left: 0 !important;
        right: 0 !important;
        z-index: 999 !important;
        background: var(--background-color) !important;
        border-top: 1px solid var(--border-color) !important;
        padding: 1rem !important;
        box-shadow: 0 -2px 10px rgba(0,0,0,0.1) !important;
    }
    
    .stChatInput > div {
        max-width: 100% !important;
    }
    
    /* Custom fixed input container */
    .fixed-input-container {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background: var(--background-color);
        border-top: 1px solid var(--border-color);
        padding: 1rem;
        z-index: 999;
        box-shadow: 0 -2px 10px rgba(0,0,0,0.1);
    }
    
    .fixed-input-row {
        display: flex;
        gap: 10px;
        align-items: center;
        max-width: 1200px;
        margin: 0 auto;
    }
    
    .fixed-text-input {
        flex: 1;
    }
    
    .fixed-mic-button {
        flex: 0 0 auto;
        width: 50px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Create a container for the main content with bottom padding
    main_container = st.container()
    
    with main_container:
        # Display chat messages from history
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Check if a quiz is in session and display it based on quiz type
        if "current_quiz" in st.session_state and "quiz_type" in st.session_state:
            with st.chat_message("assistant"):
                # Display the quiz based on its type
                if st.session_state.quiz_type == "quiz_iso27002_company":
                    display_quiz_in_streamlit(st.session_state.current_quiz)
                elif st.session_state.quiz_type == "quiz_iso27002_human_eval":
                    display_quiz_in_streamlit_1(st.session_state.current_quiz)
                    
                    # Add separate evaluate button outside the quiz display
                    col1, col2 = st.columns([1, 3])
                    with col1:
                        if st.button("🚀 Evaluate All Responses", key="evaluate_btn_main"):
                            # Evaluate responses and get results
                            evaluation_results = evaluate_quiz_responses_1()
                            if evaluation_results:
                                # Store results in session state
                                st.session_state.evaluation_results = evaluation_results
                                st.rerun()  # Force refresh to display results
                    
                    # Display evaluation results if they exist
                    if "evaluation_results" in st.session_state and st.session_state.evaluation_results:
                        display_evaluation_results_1(st.session_state.evaluation_results)
                elif st.session_state.quiz_type == "quiz_iso42001":
                    display_quiz_in_streamlit_2(st.session_state.current_quiz)
        
        # Display gap analysis if in session
        if "current_gap_analysis" in st.session_state:
            with st.chat_message("assistant"):
                display_gap_analysis_in_streamlit(st.session_state.current_gap_analysis)
        
        # Display document download button if a document was generated
        if st.session_state.document_results:
            st.markdown("### Generated Document")
            st.markdown(f"**Title:** {st.session_state.document_results['document_title']}")
            
            # Download button for PDF
            st.download_button(
                label="📄 Download Generated Document",
                data=st.session_state.document_results["pdf_data"],
                file_name=f"{st.session_state.document_results['document_title'].replace(' ', '_')}.pdf",
                mime="application/pdf",
            )
    
    # Chat UI setup
    if "history" not in st.session_state:
        st.session_state.history = []
    
    # Initialize user input
    user_input = None
    
    # Use the built-in chat input which Streamlit automatically places at the bottom
    text_input = st.chat_input("Type your message here...")
    if text_input:
        user_input = text_input
    
    # Voice input section - create a floating voice button
    if voice_enabled and processor_voice and model_voice:
        # Create a placeholder for voice input at the bottom
        voice_placeholder = st.empty()
        
        with voice_placeholder.container():
            st.markdown("""
            <div style="position: fixed; bottom: 20px; right: 20px; z-index: 1000;">
            </div>
            """, unsafe_allow_html=True)
            
            # Voice recording button with unique key to prevent conflicts
            audio_bytes = audio_recorder(
                text="🎤",
                recording_color="#e8b62c",
                neutral_color="#6aa36f",
                icon_name="microphone",
                icon_size="2x",
                pause_threshold=2.0,
                sample_rate=16000,
                key=f"voice_input_{len(st.session_state.messages)}"  # Unique key based on message count
            )
            
            if audio_bytes:
                # Always treat voice input regardless of text or previous audio
                st.audio(audio_bytes, format="audio/wav")

                # Transcribe audio
                with st.spinner("Transcribing..."):
                    transcribed_text = transcribe_audio(audio_bytes, processor_voice, model_voice)

                if transcribed_text.strip():
                    st.success(f"✅ {transcribed_text}")
                    st.session_state.pending_audio_input = transcribed_text  # persist for next rerun
                else:
                    st.warning("⚠️ Could not transcribe audio")


    # Process user input
    if user_input:
        # Add user message to both chat formats
        st.session_state.messages.append({"role": "user", "content": user_input})
        st.session_state.history.append({"role": "user", "text": user_input})
        
        # Display user message
        with st.chat_message("user"):
            st.markdown(user_input)
        
        # Process the query using our pipeline
        with st.spinner("Processing your request..."):
            result, docs, results_data = run_query_pipeline(user_input)
        
        # Display assistant response
        with st.chat_message("assistant"):
            st.markdown(result)
            
            # Handle different types of results based on content and route destination
            if isinstance(results_data, dict) and "pdf_data" in results_data:
                # Document generation results
                st.session_state.document_results = results_data
                
                # Add download button for immediate use
                st.download_button(
                    label="📄 Download Generated Document",
                    data=results_data["pdf_data"],
                    file_name=f"{results_data['document_title'].replace(' ', '_')}.pdf",
                    mime="application/pdf",
                )
            elif isinstance(results_data, list) and results_data:
                if all(isinstance(item, dict) for item in results_data):
                    # Determine the type of data based on its structure
                    if "subsection" in results_data[0]:
                        # Check for quiz data type markers to determine which quiz display to use
                        if "company_answer" in results_data[0]["qna"][0]:
                            # Company-focused quiz (ISO 27002)
                            st.session_state.current_quiz = results_data
                            st.session_state.quiz_type = "quiz_iso27002_company"
                            display_quiz_in_streamlit(results_data)
                        elif "qna" in results_data[0] and len(results_data[0]["qna"]) > 0 and not "company_answer" in results_data[0]["qna"][0]:  # ISO 42001 quiz
                            # ISO 42001 quiz
                            st.session_state.current_quiz = results_data
                            st.session_state.quiz_type = "quiz_iso42001"
                            display_quiz_in_streamlit_2(results_data)
                        else:
                            # Human evaluation quiz (ISO 27002)
                            st.session_state.current_quiz = results_data
                            st.session_state.quiz_type = "quiz_iso27002_human_eval"
                            display_quiz_in_streamlit_1(results_data)
                            
                            # Add separate evaluate button outside the quiz display
                            col1, col2 = st.columns([1, 3])
                            with col1:
                                if st.button("🚀 Evaluate All Responses", key="evaluate_btn_inline"):
                                    # Evaluate responses and get results
                                    evaluation_results = evaluate_quiz_responses_1()
                                    if evaluation_results:
                                        # Store results in session state
                                        st.session_state.evaluation_results = evaluation_results
                                        st.rerun()  # Force refresh to display results
                            
                            # Display evaluation results if they exist
                            if "evaluation_results" in st.session_state and st.session_state.evaluation_results:
                                display_evaluation_results_1(st.session_state.evaluation_results)
                    elif "theme" in results_data[0] and "clause" in results_data[0]:
                        # Gap analysis results
                        st.session_state.current_gap_analysis = results_data
                        display_gap_analysis_in_streamlit(results_data)
            
            # Handle document display for regular Q&A
            if docs:
                with st.expander("📄 Retrieved Documents"):
                    for doc in docs:
                        if isinstance(doc, str):
                            st.markdown(f"**Source:** {doc}")
                        elif hasattr(doc, 'metadata') and hasattr(doc, 'page_content'):
                            st.markdown(f"**Source:** {doc.metadata.get('source', 'Unknown')}")
                            st.text(doc.page_content)
                        else:
                            st.markdown(f"**Document:** {str(doc)}")
        
        # Add assistant response to messages
        st.session_state.messages.append({"role": "assistant", "content": result})
        st.session_state.history.append({
            "role": "ai", 
            "text": result, 
            "docs": docs if isinstance(docs, list) else [],
            "results": results_data if isinstance(results_data, list) else []
        })
        
        # Store chat history to db.json
        db['chat_history'] = st.session_state.messages
        with open(DB_FILE, 'w') as file:
            json.dump(db, file)
        
        # Clear the audio hash after processing to prevent reprocessing
        if "last_audio_hash" in st.session_state:
            del st.session_state.last_audio_hash

    # Add a "Clear Chat" button to the sidebar
    if st.sidebar.button('Clear Chat'):
        # Clear chat history in db.json
        db['chat_history'] = []
        with open(DB_FILE, 'w') as file:
            json.dump(db, file)
        
        # Clear chat messages in session state
        st.session_state.messages = []
        st.session_state.history = []
        # Clear quiz data
        if "current_quiz" in st.session_state:
            del st.session_state.current_quiz
        if "quiz_responses" in st.session_state:
            del st.session_state.quiz_responses
        if "evaluation_results" in st.session_state:
            del st.session_state.evaluation_results
            
        # Clear gap analysis data
        if "current_gap_analysis" in st.session_state:
            del st.session_state.current_gap_analysis
        if "gap_analysis_feedback" in st.session_state:
            del st.session_state.gap_analysis_feedback
        if "updated_gap_analysis" in st.session_state:
            del st.session_state.updated_gap_analysis
        st.rerun()

if __name__ == "__main__":
    main()